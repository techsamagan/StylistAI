"""
Appends Appendix B (full source code), Appendix C (DB schema), and
Appendix D (frontend architecture) to the existing StylistAI report.
Run AFTER generate_report.py has already produced the .docx file.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = "/Users/samagannurdinov/Desktop/StylistAI/final/StylistAI_Final_Report_Samagan_Nurdinov.docx"
doc  = Document(DOCX)

# ── Helpers ────────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, before=24, after=8, size=None, color=None):
    sz = (size or 16) if level == 1 else (size or 14) if level == 2 else (size or 12)
    p  = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    for run in p.runs:
        run.font.name      = "Times New Roman"
        run.font.size      = Pt(sz)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(*(color or (0, 0, 0)))
    return p

def body(text, before=0, after=6, indent=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = Pt(22)
    if indent:
        p.paragraph_format.first_line_indent = Inches(indent)
    r = p.add_run(text)
    set_font(r, size=12, bold=bold)
    return p

def para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0,
         size=12, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p

def bullet(text, before=2, after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    set_font(r, size=12)
    return p

def cb(text):
    """code_block"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    return p

def pb():
    doc.add_page_break()

def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
    return row

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX B — COMPLETE SOURCE CODE LISTINGS
# ══════════════════════════════════════════════════════════════════════════════
pb()
heading("Appendix B: Complete Source Code Listings", 1)
body(
    "This appendix presents the complete, annotated source code for all major modules "
    "of the StylistAI application. The listings are organised into backend modules "
    "(Python / FastAPI) followed by frontend modules (React / JavaScript). Each listing "
    "is preceded by a brief description of the module's role and its key design decisions. "
    "The code is presented exactly as it exists in the final submitted version of the "
    "project, including all imports, type annotations, and dependency-injection patterns.",
    indent=0.5
)
body(
    "All backend source files reside under the style_back/ directory. The entry point is "
    "style_back/main.py, which wires together all routers and configures the ASGI "
    "application. Database models live in style_back/app/models.py, Pydantic schemas in "
    "style_back/app/schemas.py, and the eight router modules in style_back/app/routers/. "
    "The frontend source resides under style_front/src/, with the React application root "
    "at style_front/src/App.jsx and all HTTP communication encapsulated in "
    "style_front/src/api/client.js.",
    indent=0.5
)

# ── B.1 ──────────────────────────────────────────────────────────────────────
heading("B.1  Backend Entry Point: style_back/main.py", 2)
body(
    "The main.py file is the ASGI application entry point. It instantiates the FastAPI "
    "application object, registers the CORS middleware to permit cross-origin requests from "
    "the React development server running on port 3000, mounts all eight router modules, "
    "and serves uploaded garment images as static files under the /uploads path. The "
    "Base.metadata.create_all() call at module load time ensures that all SQLAlchemy ORM "
    "models are reflected into the SQLite database before the first HTTP request is "
    "processed, effectively auto-migrating the schema on application startup.",
    indent=0.5
)
cb(
    "from pathlib import Path\n"
    "from fastapi import FastAPI\n"
    "from fastapi.middleware.cors import CORSMiddleware\n"
    "from fastapi.staticfiles import StaticFiles\n"
    "\n"
    "from app.database import engine\n"
    "from app.models import Base\n"
    "from app.routers import (health, closet, outfits, auth,\n"
    "                          suggestions, calendar, weather, travel)\n"
    "\n"
    "Base.metadata.create_all(bind=engine)\n"
    "\n"
    "UPLOADS_DIR = Path(__file__).parent / \"uploads\"\n"
    "UPLOADS_DIR.mkdir(exist_ok=True)\n"
    "\n"
    "app = FastAPI(\n"
    "    title=\"Style API\",\n"
    "    description=\"Backend for AI-powered wardrobe and outfit suggestions\",\n"
    "    version=\"0.1.0\",\n"
    ")\n"
    "\n"
    "app.add_middleware(\n"
    "    CORSMiddleware,\n"
    "    allow_origins=[\"http://localhost:3000\", \"http://127.0.0.1:3000\"],\n"
    "    allow_credentials=True,\n"
    "    allow_methods=[\"*\"],\n"
    "    allow_headers=[\"*\"],\n"
    ")"
)
cb(
    "app.include_router(health.router)\n"
    "app.include_router(closet.router)\n"
    "app.include_router(outfits.router)\n"
    "app.include_router(auth.router)\n"
    "app.include_router(suggestions.router)\n"
    "app.include_router(calendar.router)\n"
    "app.include_router(weather.router)\n"
    "app.include_router(travel.router)\n"
    "\n"
    "app.mount(\"/uploads\",\n"
    "          StaticFiles(directory=str(UPLOADS_DIR)), name=\"uploads\")\n"
    "\n"
    "\n"
    "@app.get(\"/\")\n"
    "def root():\n"
    "    return {\"message\": \"Style API\", \"docs\": \"/docs\"}"
)

# ── B.2 ──────────────────────────────────────────────────────────────────────
heading("B.2  Database Configuration: style_back/app/database.py", 2)
body(
    "The database.py module centralises all SQLAlchemy engine and session configuration. "
    "The DATABASE_URL environment variable is read at startup from a .env file via "
    "python-dotenv. When the URL starts with 'sqlite', the engine is created with the "
    "check_same_thread=False argument, which is required for SQLite to operate safely "
    "within FastAPI's async request handling (which can dispatch handlers on different "
    "threads). The pool_pre_ping=True flag causes SQLAlchemy to test every connection "
    "before handing it out, preventing stale-connection errors after idle periods. The "
    "get_db() generator function implements the per-request database session lifecycle "
    "used throughout all routers via FastAPI's Depends() dependency injection mechanism.",
    indent=0.5
)
cb(
    "import os\n"
    "from dotenv import load_dotenv\n"
    "from sqlalchemy import create_engine\n"
    "from sqlalchemy.orm import sessionmaker, DeclarativeBase\n"
    "\n"
    "load_dotenv()\n"
    "\n"
    "DATABASE_URL = os.getenv(\"DATABASE_URL\", \"\")\n"
    "\n"
    "if DATABASE_URL.startswith(\"sqlite\"):\n"
    "    engine = create_engine(\n"
    "        DATABASE_URL,\n"
    "        connect_args={\"check_same_thread\": False},\n"
    "        pool_pre_ping=True\n"
    "    )\n"
    "else:\n"
    "    engine = create_engine(DATABASE_URL, pool_pre_ping=True)\n"
    "\n"
    "SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)\n"
    "\n"
    "\n"
    "class Base(DeclarativeBase):\n"
    "    pass\n"
    "\n"
    "\n"
    "def get_db():\n"
    "    db = SessionLocal()\n"
    "    try:\n"
    "        yield db\n"
    "    finally:\n"
    "        db.close()"
)

# ── B.3 ──────────────────────────────────────────────────────────────────────
heading("B.3  Pydantic Schemas: style_back/app/schemas.py", 2)
body(
    "The schemas.py module defines all Pydantic v2 data models used for request "
    "validation, response serialisation, and OpenAPI documentation generation. FastAPI "
    "uses these models to automatically validate incoming JSON payloads and to serialise "
    "SQLAlchemy ORM objects into JSON responses. The Category enum ensures that only the "
    "five defined clothing categories are accepted at the API boundary, providing "
    "early-failure validation. The ClosetItem model includes the from_attributes = True "
    "configuration (Pydantic v2 equivalent of orm_mode) to enable direct serialisation "
    "from SQLAlchemy model instances. The OutfitRequest model uses Field with ge/le "
    "constraints to enforce the 0-100 range of the vibe parameter at the API boundary "
    "without requiring manual validation in the route handler.",
    indent=0.5
)
cb(
    "from pydantic import BaseModel, Field\n"
    "from typing import Optional\n"
    "from enum import Enum\n"
    "\n"
    "\n"
    "class Category(str, Enum):\n"
    "    TOP       = \"Top\"\n"
    "    BOTTOM    = \"Bottom\"\n"
    "    OUTERWEAR = \"Outerwear\"\n"
    "    SHOES     = \"Shoes\"\n"
    "    ACCESSORY = \"Accessory\"\n"
    "\n"
    "\n"
    "class ClosetItemCreate(BaseModel):\n"
    "    name:            str\n"
    "    category:        Category\n"
    "    image_url:       Optional[str] = None\n"
    "    color:           Optional[str] = None\n"
    "    formality:       Optional[str] = None   # CASUAL | MODERATE | FORMAL | UNIVERSAL\n"
    "    formality_value: Optional[int] = None   # 0-100 for progress bar display\n"
    "\n"
    "\n"
    "class ClosetItem(ClosetItemCreate):\n"
    "    id:      int\n"
    "    user_id: Optional[str] = None\n"
    "\n"
    "    class Config:\n"
    "        from_attributes = True\n"
    "\n"
    "\n"
    "class ClosetItemUpdate(BaseModel):\n"
    "    name:            Optional[str]      = None\n"
    "    category:        Optional[Category] = None\n"
    "    image_url:       Optional[str]      = None\n"
    "    color:           Optional[str]      = None\n"
    "    formality:       Optional[str]      = None\n"
    "    formality_value: Optional[int]      = None"
)
cb(
    "class OutfitRequest(BaseModel):\n"
    "    context:              str            = Field(..., description=\"Office, Date Night, Travel, Gym\")\n"
    "    weather_temp_c:       Optional[float] = None\n"
    "    formality_preference: Optional[str]  = None\n"
    "    vibe:                 Optional[int]  = Field(None, ge=0, le=100,\n"
    "                                                  description=\"0=comfort, 100=style\")\n"
    "\n"
    "\n"
    "class OutfitItem(BaseModel):\n"
    "    id:        int\n"
    "    name:      str\n"
    "    category:  str\n"
    "    image_url: Optional[str] = None\n"
    "\n"
    "\n"
    "class OutfitSuggestion(BaseModel):\n"
    "    items:       list[OutfitItem]\n"
    "    explanation: str\n"
    "\n"
    "\n"
    "class SaveOutfitRequest(BaseModel):\n"
    "    context:     str\n"
    "    items:       list[OutfitItem]\n"
    "    explanation: str\n"
    "\n"
    "\n"
    "class SavedOutfit(BaseModel):\n"
    "    id:          str\n"
    "    context:     str\n"
    "    items:       list[OutfitItem]\n"
    "    explanation: str\n"
    "    saved_at:    str"
)
cb(
    "class UserCreate(BaseModel):\n"
    "    email:    str\n"
    "    password: str\n"
    "    name:     Optional[str] = None\n"
    "    city:     Optional[str] = None   # e.g. 'London' used for live weather\n"
    "\n"
    "\n"
    "class UserUpdate(BaseModel):\n"
    "    name:     Optional[str] = None\n"
    "    city:     Optional[str] = None\n"
    "    email:    Optional[str] = None\n"
    "    password: Optional[str] = None\n"
    "\n"
    "\n"
    "class UserProfile(BaseModel):\n"
    "    email: str\n"
    "    name:  Optional[str] = None\n"
    "    city:  Optional[str] = None\n"
    "\n"
    "\n"
    "class UserLogin(BaseModel):\n"
    "    email:    str\n"
    "    password: str\n"
    "\n"
    "\n"
    "class Token(BaseModel):\n"
    "    access_token: str\n"
    "    token_type:   str           = \"bearer\"\n"
    "    city:         Optional[str] = None\n"
    "    name:         Optional[str] = None\n"
    "\n"
    "\n"
    "class AISuggestionResponse(BaseModel):\n"
    "    item_name: str\n"
    "    reason:    str"
)
cb(
    "class CalendarEvent(BaseModel):\n"
    "    id:        str\n"
    "    title:     str\n"
    "    start:     str            # ISO datetime e.g. '2026-03-24T10:00'\n"
    "    end:       str\n"
    "    formality: Optional[str] = None\n"
    "\n"
    "\n"
    "class PackRequest(BaseModel):\n"
    "    destination: str\n"
    "    days:        int\n"
    "\n"
    "\n"
    "class DailyOutfit(BaseModel):\n"
    "    day:   int\n"
    "    items: list[OutfitItem]\n"
    "\n"
    "\n"
    "class PackResponse(BaseModel):\n"
    "    destination:   str\n"
    "    weather_summary: str\n"
    "    packing_list:  list[OutfitItem]\n"
    "    daily_outfits: list[DailyOutfit]\n"
    "\n"
    "\n"
    "class TryOnRequest(BaseModel):\n"
    "    outfit_items: list[str]\n"
    "    context:      str\n"
    "\n"
    "\n"
    "class TryOnResponse(BaseModel):\n"
    "    image_url: str"
)

# ── B.4 ──────────────────────────────────────────────────────────────────────
heading("B.4  Authentication Router: style_back/app/routers/auth.py", 2)
body(
    "The authentication router exposes four endpoints: POST /auth/register for new user "
    "creation, POST /auth/login for credential verification, GET /auth/me for profile "
    "retrieval, and PUT /auth/me for profile updates. Passwords are hashed using bcrypt "
    "with an automatically generated salt via bcrypt.hashpw(). The access token returned "
    "upon successful login or registration is a simple string in the format "
    "'token-{email}'. While this is intentionally simplified for an academic prototype, "
    "the architecture is designed so that this could be replaced with a proper JWT "
    "implementation by modifying only the get_current_user() function in auth_utils.py "
    "without changes to any of the four route handlers. The update endpoint checks for "
    "email uniqueness before committing a change, preventing two users from sharing the "
    "same email address after a profile update.",
    indent=0.5
)
cb(
    "import bcrypt\n"
    "from fastapi import APIRouter, HTTPException, Depends\n"
    "from sqlalchemy.orm import Session\n"
    "from app.schemas import UserCreate, UserLogin, Token, UserProfile, UserUpdate\n"
    "from app.database import get_db\n"
    "from app.models import UserModel\n"
    "from app.auth_utils import get_current_user\n"
    "\n"
    "router = APIRouter(prefix=\"/auth\", tags=[\"auth\"])\n"
    "\n"
    "\n"
    "def _hash(password: str) -> str:\n"
    "    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()\n"
    "\n"
    "\n"
    "def _verify(password: str, hashed: str) -> bool:\n"
    "    return bcrypt.checkpw(password.encode(), hashed.encode())\n"
    "\n"
    "\n"
    "@router.post(\"/register\", response_model=Token, status_code=201)\n"
    "def register(user: UserCreate, db: Session = Depends(get_db)):\n"
    "    if db.query(UserModel).filter(UserModel.email == user.email).first():\n"
    "        raise HTTPException(status_code=400, detail=\"User already exists\")\n"
    "    db_user = UserModel(\n"
    "        email=user.email,\n"
    "        password_hash=_hash(user.password),\n"
    "        name=user.name,\n"
    "        city=user.city,\n"
    "    )\n"
    "    db.add(db_user)\n"
    "    db.commit()\n"
    "    db.refresh(db_user)\n"
    "    return Token(access_token=f\"token-{user.email}\",\n"
    "                 city=user.city, name=user.name)"
)
cb(
    "@router.post(\"/login\", response_model=Token)\n"
    "def login(credentials: UserLogin, db: Session = Depends(get_db)):\n"
    "    user = db.query(UserModel).filter(\n"
    "               UserModel.email == credentials.email).first()\n"
    "    if not user or not _verify(credentials.password, user.password_hash):\n"
    "        raise HTTPException(status_code=400,\n"
    "                            detail=\"Invalid email or password\")\n"
    "    return Token(access_token=f\"token-{user.email}\",\n"
    "                 city=user.city, name=user.name)\n"
    "\n"
    "\n"
    "@router.get(\"/me\", response_model=UserProfile)\n"
    "def get_me(user: UserModel = Depends(get_current_user)):\n"
    "    if not user:\n"
    "        raise HTTPException(status_code=401, detail=\"Not authenticated\")\n"
    "    return UserProfile(email=user.email, name=user.name, city=user.city)\n"
    "\n"
    "\n"
    "@router.put(\"/me\", response_model=Token)\n"
    "def update_me(update: UserUpdate,\n"
    "              db:   Session   = Depends(get_db),\n"
    "              user: UserModel = Depends(get_current_user)):\n"
    "    if not user:\n"
    "        raise HTTPException(status_code=401, detail=\"Not authenticated\")\n"
    "    if update.email and update.email != user.email:\n"
    "        if db.query(UserModel).filter(\n"
    "               UserModel.email == update.email).first():\n"
    "            raise HTTPException(status_code=400,\n"
    "                                detail=\"Email already taken\")\n"
    "        user.email = update.email\n"
    "    if update.name     is not None: user.name          = update.name\n"
    "    if update.city     is not None: user.city          = update.city\n"
    "    if update.password:             user.password_hash = _hash(update.password)\n"
    "    db.commit(); db.refresh(user)\n"
    "    return Token(access_token=f\"token-{user.email}\",\n"
    "                 city=user.city, name=user.name)"
)

# ── B.5 ──────────────────────────────────────────────────────────────────────
heading("B.5  Closet Management Router: style_back/app/routers/closet.py", 2)
body(
    "The closet router manages the CRUD lifecycle of wardrobe items. Its most technically "
    "interesting endpoint is POST /closet/upload, which accepts a multipart form upload "
    "containing an image file alongside item metadata. The background removal pipeline is "
    "invoked here: the raw image bytes are passed to rembg.remove(), which applies the "
    "U2-Net salient object detection model to produce a PNG with a transparent background. "
    "This processed image is then written to the uploads directory and served as a static "
    "file, with the full URL stored against the closet item in the database. If the rembg "
    "library is unavailable or the model fails to load, the raw image bytes are saved "
    "without processing (graceful degradation). The list endpoint supports four independent "
    "query-string filters (category, color, formality, search) that are composed "
    "dynamically into a SQLAlchemy query, allowing the frontend to filter the wardrobe "
    "grid without additional API calls.",
    indent=0.5
)
cb(
    "import uuid\n"
    "from pathlib import Path\n"
    "from typing import Optional\n"
    "from fastapi import (APIRouter, HTTPException, Depends,\n"
    "                     UploadFile, File, Form, Request)\n"
    "from sqlalchemy.orm import Session\n"
    "from app.schemas import ClosetItem, ClosetItemCreate, ClosetItemUpdate, Category\n"
    "from app.database import get_db\n"
    "from app.models import ClosetItemModel\n"
    "from app.auth_utils import get_current_user_id\n"
    "\n"
    "router = APIRouter(prefix=\"/closet\", tags=[\"closet\"])\n"
    "UPLOADS_DIR = Path(__file__).parent.parent.parent / \"uploads\"\n"
    "UPLOADS_DIR.mkdir(exist_ok=True)\n"
    "\n"
    "\n"
    "def _to_schema(item: ClosetItemModel) -> ClosetItem:\n"
    "    return ClosetItem(\n"
    "        id=item.id, user_id=str(item.user_id) if item.user_id else None,\n"
    "        name=item.name, category=Category(item.category),\n"
    "        image_url=item.image_url, color=item.color,\n"
    "        formality=item.formality, formality_value=item.formality_value,\n"
    "    )"
)
cb(
    "@router.get(\"\", response_model=list[ClosetItem])\n"
    "def list_items(\n"
    "    category: Optional[Category] = None,\n"
    "    color:    Optional[str]      = None,\n"
    "    formality:Optional[str]      = None,\n"
    "    search:   Optional[str]      = None,\n"
    "    db:       Session            = Depends(get_db),\n"
    "    user_id:  Optional[int]      = Depends(get_current_user_id),\n"
    "):\n"
    "    q = db.query(ClosetItemModel)\n"
    "    if user_id:               q = q.filter(ClosetItemModel.user_id   == user_id)\n"
    "    if category:              q = q.filter(ClosetItemModel.category  == category.value)\n"
    "    if color and color.strip():\n"
    "        q = q.filter(ClosetItemModel.color == color.strip().lower())\n"
    "    if formality and formality.strip():\n"
    "        q = q.filter(ClosetItemModel.formality == formality.strip().upper())\n"
    "    if search and search.strip():\n"
    "        s = f\"%{search.strip().lower()}%\"\n"
    "        q = q.filter(\n"
    "            ClosetItemModel.name.ilike(s)     |\n"
    "            ClosetItemModel.category.ilike(s) |\n"
    "            ClosetItemModel.color.ilike(s)\n"
    "        )\n"
    "    return [_to_schema(i) for i in q.all()]"
)
cb(
    "@router.post(\"/upload\", response_model=ClosetItem, status_code=201)\n"
    "async def upload_item(\n"
    "    request:         Request,\n"
    "    name:            str            = Form(...),\n"
    "    category:        str            = Form(\"Top\"),\n"
    "    color:           Optional[str]  = Form(None),\n"
    "    formality:       str            = Form(\"MODERATE\"),\n"
    "    formality_value: int            = Form(50),\n"
    "    file:            Optional[UploadFile] = File(None),\n"
    "    db:              Session        = Depends(get_db),\n"
    "    user_id:         Optional[int]  = Depends(get_current_user_id),\n"
    "):\n"
    "    image_url = None\n"
    "    if file and file.filename:\n"
    "        raw_bytes = await file.read()\n"
    "        try:\n"
    "            from rembg import remove\n"
    "            output_bytes = remove(raw_bytes)\n"
    "            filename = f\"{uuid.uuid4()}.png\"\n"
    "            (UPLOADS_DIR / filename).write_bytes(output_bytes)\n"
    "        except Exception:\n"
    "            ext = Path(file.filename).suffix.lower() or \".jpg\"\n"
    "            filename = f\"{uuid.uuid4()}{ext}\"\n"
    "            (UPLOADS_DIR / filename).write_bytes(raw_bytes)\n"
    "        base = str(request.base_url).rstrip(\"/\")\n"
    "        image_url = f\"{base}/uploads/{filename}\"\n"
    "    db_item = ClosetItemModel(\n"
    "        user_id=user_id, name=name, category=category,\n"
    "        image_url=image_url, color=color or None,\n"
    "        formality=formality, formality_value=formality_value,\n"
    "    )\n"
    "    db.add(db_item); db.commit(); db.refresh(db_item)\n"
    "    return _to_schema(db_item)"
)
cb(
    "@router.get(\"/{item_id}\", response_model=ClosetItem)\n"
    "def get_item(item_id: int, db: Session = Depends(get_db),\n"
    "             user_id: Optional[int] = Depends(get_current_user_id)):\n"
    "    item = db.query(ClosetItemModel).filter(\n"
    "               ClosetItemModel.id == item_id).first()\n"
    "    if not item:\n"
    "        raise HTTPException(status_code=404, detail=\"Item not found\")\n"
    "    return _to_schema(item)\n"
    "\n"
    "\n"
    "@router.patch(\"/{item_id}\", response_model=ClosetItem)\n"
    "def update_item(item_id: int, update: ClosetItemUpdate,\n"
    "                db: Session = Depends(get_db),\n"
    "                user_id: Optional[int] = Depends(get_current_user_id)):\n"
    "    item = db.query(ClosetItemModel).filter(\n"
    "               ClosetItemModel.id == item_id).first()\n"
    "    if not item:\n"
    "        raise HTTPException(status_code=404, detail=\"Item not found\")\n"
    "    for field, value in update.model_dump(exclude_unset=True).items():\n"
    "        setattr(item, field,\n"
    "                value if not hasattr(value, 'value') else value.value)\n"
    "    db.commit(); db.refresh(item)\n"
    "    return _to_schema(item)\n"
    "\n"
    "\n"
    "@router.delete(\"/{item_id}\", status_code=204)\n"
    "def delete_item(item_id: int, db: Session = Depends(get_db),\n"
    "                user_id: Optional[int] = Depends(get_current_user_id)):\n"
    "    item = db.query(ClosetItemModel).filter(\n"
    "               ClosetItemModel.id == item_id).first()\n"
    "    if not item:\n"
    "        raise HTTPException(status_code=404, detail=\"Item not found\")\n"
    "    db.delete(item); db.commit()\n"
    "    return None"
)

# ── B.6 ──────────────────────────────────────────────────────────────────────
heading("B.6  Outfit Generation Router: style_back/app/routers/outfits.py", 2)
body(
    "The outfits router is the most complex module in the backend and encapsulates three "
    "distinct outfit-related features: rule-based outfit generation (POST /outfits/generate), "
    "contextual daily outfit recommendation integrating live weather and calendar data "
    "(GET /outfits/today), outfit persistence (POST /outfits/save, GET /outfits/saved, "
    "DELETE /outfits/saved/{id}), and AI-powered virtual try-on via DALL-E 3 "
    "(POST /outfits/try-on). The generate endpoint uses a vibe score (0-100 integer) to "
    "determine the formality target set, then applies the _pick() helper to select one "
    "item per clothing category from the filtered pool. The today endpoint additionally "
    "fetches real-time weather from Open-Meteo, queries the user's calendar events for "
    "the current day, and uses a deterministic random seed derived from the date to "
    "ensure the same outfit is served on repeated requests within the same day, unless "
    "the ?refresh=true query parameter is passed to request a different combination.",
    indent=0.5
)
cb(
    "import json, uuid, random, datetime, urllib.parse, os\n"
    "from typing import Optional\n"
    "from fastapi import APIRouter, HTTPException, Depends\n"
    "from sqlalchemy.orm import Session\n"
    "from openai import OpenAI\n"
    "from app.schemas import (OutfitRequest, OutfitSuggestion, OutfitItem,\n"
    "                          SaveOutfitRequest, SavedOutfit,\n"
    "                          TryOnRequest, TryOnResponse)\n"
    "from app.database import get_db\n"
    "from app.models import (ClosetItemModel, CalendarEventModel,\n"
    "                         SavedOutfitModel, UserModel)\n"
    "from app.auth_utils import get_current_user_id, get_current_user\n"
    "from app.routers.weather import _fetch_json\n"
    "\n"
    "_FORMALITY_RANK = {\"formal\": 3, \"business_casual\": 2, \"casual\": 1}\n"
    "_RAIN_CODES     = {51, 53, 55, 61, 63, 65, 80, 81, 82, 95, 96, 99}\n"
    "\n"
    "router = APIRouter(prefix=\"/outfits\", tags=[\"outfits\"])\n"
    "\n"
    "\n"
    "def _pick(items, category, formalities):\n"
    "    pool = [i for i in items\n"
    "            if i.category == category and i.formality in formalities]\n"
    "    return pool or [i for i in items if i.category == category]"
)
cb(
    "def _get_weather(city: str):\n"
    "    try:\n"
    "        geo = _fetch_json(\n"
    "            f\"https://geocoding-api.open-meteo.com/v1/search\"\n"
    "            f\"?name={urllib.parse.quote(city)}&count=1&language=en&format=json\"\n"
    "        )\n"
    "        results = geo.get(\"results\") or []\n"
    "        if not results: return None\n"
    "        r  = results[0]\n"
    "        wx = _fetch_json(\n"
    "            f\"https://api.open-meteo.com/v1/forecast\"\n"
    "            f\"?latitude={r['latitude']}&longitude={r['longitude']}\"\n"
    "            f\"&current=temperature_2m,weathercode\"\n"
    "            f\"&temperature_unit=celsius&timezone=auto\"\n"
    "        )\n"
    "        current = wx.get(\"current\", {})\n"
    "        return {\n"
    "            \"temp_c\": current.get(\"temperature_2m\"),\n"
    "            \"code\":   current.get(\"weathercode\", 0),\n"
    "            \"city\":   r.get(\"name\", city),\n"
    "        }\n"
    "    except Exception:\n"
    "        return None"
)
cb(
    "@router.post(\"/generate\", response_model=OutfitSuggestion)\n"
    "def generate_outfit(\n"
    "    req:     OutfitRequest,\n"
    "    db:      Session        = Depends(get_db),\n"
    "    user_id: Optional[int]  = Depends(get_current_user_id),\n"
    "):\n"
    "    q = db.query(ClosetItemModel)\n"
    "    if user_id:\n"
    "        q = q.filter(ClosetItemModel.user_id == user_id)\n"
    "    items = q.all()\n"
    "\n"
    "    vibe = req.vibe if req.vibe is not None else 50\n"
    "    if   vibe < 30: target = {\"CASUAL\", \"MODERATE\"}\n"
    "    elif vibe > 70: target = {\"FORMAL\", \"MODERATE\"}\n"
    "    else:           target = {\"CASUAL\", \"MODERATE\", \"FORMAL\", \"UNIVERSAL\"}\n"
    "\n"
    "    tops    = _pick(items, \"Top\",       target)\n"
    "    bottoms = _pick(items, \"Bottom\",    target)\n"
    "    shoes   = _pick(items, \"Shoes\",     target)\n"
    "    outer   = _pick(items, \"Outerwear\", target)\n"
    "\n"
    "    picked = []\n"
    "    for item in [tops[0] if tops else None,\n"
    "                 bottoms[0] if bottoms else None,\n"
    "                 shoes[0]   if shoes   else None]:\n"
    "        if item:\n"
    "            picked.append(OutfitItem(id=item.id, name=item.name,\n"
    "                                     category=item.category,\n"
    "                                     image_url=item.image_url))\n"
    "    if outer and (req.weather_temp_c is None or req.weather_temp_c < 18):\n"
    "        picked.append(OutfitItem(id=outer[0].id, name=outer[0].name,\n"
    "                                 category=outer[0].category,\n"
    "                                 image_url=outer[0].image_url))\n"
    "\n"
    "    vibe_note = \" Styled for comfort.\" if vibe < 30 else \\\n"
    "               \" Styled for impact.\"  if vibe > 70 else \"\"\n"
    "    temp_note = (f\" {req.weather_temp_c:.0f}\\u00b0C \\u2014 layering recommended.\"\n"
    "                 if req.weather_temp_c is not None and req.weather_temp_c < 18\n"
    "                 else \"\")\n"
    "    explanation = (f\"This works for {req.context}: neutral base keeps it \"\n"
    "                   f\"appropriate.{vibe_note}{temp_note}\")\n"
    "    return OutfitSuggestion(items=picked, explanation=explanation)"
)
cb(
    "@router.get(\"/today\", response_model=OutfitSuggestion)\n"
    "def get_today_outfit(\n"
    "    refresh: bool                  = False,\n"
    "    db:      Session               = Depends(get_db),\n"
    "    user:    Optional[UserModel]   = Depends(get_current_user),\n"
    "):\n"
    "    user_id = user.id   if user else None\n"
    "    city    = user.city if user else None\n"
    "    q = db.query(ClosetItemModel)\n"
    "    if user_id: q = q.filter(ClosetItemModel.user_id == user_id)\n"
    "    items = q.all()\n"
    "    if not items:\n"
    "        return OutfitSuggestion(items=[],\n"
    "                                explanation=\"Add items to your closet.\")\n"
    "\n"
    "    weather  = _get_weather(city) if city else None\n"
    "    temp_c   = weather[\"temp_c\"]  if weather else None\n"
    "    is_rainy = weather and weather[\"code\"] in _RAIN_CODES\n"
    "    is_cold  = temp_c is not None and temp_c < 15\n"
    "    is_hot   = temp_c is not None and temp_c > 25\n"
    "    weather_city = weather[\"city\"] if weather else city\n"
    "\n"
    "    today  = datetime.date.today().isoformat()\n"
    "    eq = db.query(CalendarEventModel).filter(\n"
    "             CalendarEventModel.start.like(f\"{today}%\"))\n"
    "    if user_id: eq = eq.filter(CalendarEventModel.user_id == user_id)\n"
    "    events    = eq.all()\n"
    "    top_event = max(events,\n"
    "                    key=lambda e: _FORMALITY_RANK.get(e.formality or \"\", 0)\n"
    "                   ) if events else None\n"
    "\n"
    "    if top_event and top_event.formality == \"formal\":\n"
    "        target = {\"FORMAL\", \"MODERATE\"}\n"
    "    elif top_event and top_event.formality == \"business_casual\":\n"
    "        target = {\"MODERATE\", \"FORMAL\", \"UNIVERSAL\"}\n"
    "    else:\n"
    "        target = {\"CASUAL\", \"MODERATE\", \"UNIVERSAL\"}"
)
cb(
    "    seed = (int(datetime.datetime.now().timestamp()) if refresh\n"
    "            else int(datetime.date.today().strftime(\"%Y%m%d\")))\n"
    "    rng = random.Random(seed)\n"
    "\n"
    "    def pick(cat):\n"
    "        pool = ([i for i in items if i.category == cat\n"
    "                 and i.formality in target]\n"
    "                or [i for i in items if i.category == cat])\n"
    "        if not pool: return None\n"
    "        rng.shuffle(pool); return pool[0]\n"
    "\n"
    "    top    = pick(\"Top\")\n"
    "    bottom = pick(\"Bottom\")\n"
    "    shoes  = pick(\"Shoes\")\n"
    "    outer  = pick(\"Outerwear\") if (is_cold or is_rainy) else None\n"
    "\n"
    "    picked = [\n"
    "        OutfitItem(id=i.id, name=i.name,\n"
    "                   category=i.category, image_url=i.image_url)\n"
    "        for i in [top, bottom, shoes, outer] if i\n"
    "    ]\n"
    "\n"
    "    parts = []\n"
    "    if top_event: parts.append(f\"your {top_event.title} today\")\n"
    "    if temp_c is not None:\n"
    "        if is_rainy: parts.append(f\"rain expected in {weather_city} ({temp_c:.0f}\\u00b0C)\")\n"
    "        elif is_cold: parts.append(f\"cold weather in {weather_city} ({temp_c:.0f}\\u00b0C)\")\n"
    "        elif is_hot:  parts.append(f\"warm weather in {weather_city} ({temp_c:.0f}\\u00b0C)\")\n"
    "        else:         parts.append(f\"mild {temp_c:.0f}\\u00b0C in {weather_city}\")\n"
    "\n"
    "    explanation = (f\"Chosen for {' and '.join(parts)}.\" if parts\n"
    "                   else \"A well-balanced outfit for today.\")\n"
    "    return OutfitSuggestion(items=picked, explanation=explanation)"
)
cb(
    "@router.post(\"/save\", response_model=SavedOutfit, status_code=201)\n"
    "def save_outfit(\n"
    "    req:     SaveOutfitRequest,\n"
    "    db:      Session       = Depends(get_db),\n"
    "    user_id: Optional[int] = Depends(get_current_user_id),\n"
    "):\n"
    "    db_outfit = SavedOutfitModel(\n"
    "        id=str(uuid.uuid4()), user_id=user_id,\n"
    "        context=req.context,\n"
    "        items_json=json.dumps([i.model_dump() for i in req.items]),\n"
    "        explanation=req.explanation,\n"
    "    )\n"
    "    db.add(db_outfit); db.commit(); db.refresh(db_outfit)\n"
    "    return SavedOutfit(\n"
    "        id=db_outfit.id, context=db_outfit.context,\n"
    "        items=[OutfitItem(**i)\n"
    "               for i in json.loads(db_outfit.items_json)],\n"
    "        explanation=db_outfit.explanation,\n"
    "        saved_at=db_outfit.saved_at.isoformat(),\n"
    "    )\n"
    "\n"
    "\n"
    "@router.get(\"/saved\", response_model=list[SavedOutfit])\n"
    "def list_saved(db: Session = Depends(get_db),\n"
    "               user_id: Optional[int] = Depends(get_current_user_id)):\n"
    "    q = db.query(SavedOutfitModel)\n"
    "    if user_id: q = q.filter(SavedOutfitModel.user_id == user_id)\n"
    "    return [\n"
    "        SavedOutfit(\n"
    "            id=o.id, context=o.context,\n"
    "            items=[OutfitItem(**i)\n"
    "                   for i in json.loads(o.items_json)],\n"
    "            explanation=o.explanation,\n"
    "            saved_at=o.saved_at.isoformat(),\n"
    "        )\n"
    "        for o in q.all()\n"
    "    ]"
)
cb(
    "@router.post(\"/try-on\", response_model=TryOnResponse)\n"
    "def virtual_try_on(req: TryOnRequest):\n"
    "    api_key = os.getenv(\"OPENAI_API_KEY\")\n"
    "    if not api_key:\n"
    "        raise HTTPException(status_code=500,\n"
    "                            detail=\"OpenAI API key not configured\")\n"
    "    client     = OpenAI(api_key=api_key)\n"
    "    items_desc = \", \".join(req.outfit_items)\n"
    "    prompt = (\n"
    "        f\"A photorealistic, high-fashion full-body portrait of a \"\n"
    "        f\"stylish person wearing: {items_desc}. \"\n"
    "        f\"Context: {req.context}. \"\n"
    "        f\"Studio lighting, highly detailed, premium look, \"\n"
    "        f\"clear facial features, modern aesthetic.\"\n"
    "    )\n"
    "    try:\n"
    "        response = client.images.generate(\n"
    "            model=\"dall-e-3\",\n"
    "            prompt=prompt,\n"
    "            size=\"1024x1024\",\n"
    "            quality=\"standard\",\n"
    "            n=1,\n"
    "        )\n"
    "        return TryOnResponse(image_url=response.data[0].url)\n"
    "    except Exception as e:\n"
    "        raise HTTPException(status_code=500,\n"
    "                            detail=f\"Image generation failed: {str(e)}\")"
)

# ── B.7 ──────────────────────────────────────────────────────────────────────
heading("B.7  Smart Suggestions Router: style_back/app/routers/suggestions.py", 2)
body(
    "The suggestions router provides two analytical endpoints. The GET /suggestions/today "
    "endpoint recommends a single featured item from the user's wardrobe, explaining the "
    "recommendation in natural language by combining calendar event context and live "
    "weather conditions. The logic prioritises outerwear on cold or rainy days, light "
    "tops on warm days, and items whose formality level matches the highest-priority "
    "calendar event of the day. The GET /suggestions/gaps endpoint performs a wardrobe "
    "gap analysis by examining the category distribution and formality balance of the "
    "user's entire closet. It returns a single actionable shopping recommendation: if the "
    "user has many tops but no bottoms, for example, the system recommends versatile "
    "trousers and explains that this would unlock multiple new outfit combinations. This "
    "deterministic analysis requires no external API calls and executes in microseconds.",
    indent=0.5
)
cb(
    "import datetime, urllib.parse\n"
    "from typing import Optional\n"
    "from fastapi import APIRouter, Depends\n"
    "from sqlalchemy.orm import Session\n"
    "from app.schemas import AISuggestionResponse\n"
    "from app.database import get_db\n"
    "from app.models import ClosetItemModel, CalendarEventModel\n"
    "from app.auth_utils import get_current_user_id\n"
    "from app.routers.weather import _fetch_json\n"
    "\n"
    "router = APIRouter(prefix=\"/suggestions\", tags=[\"suggestions\"])\n"
    "_FORMALITY_RANK = {\"formal\": 3, \"business_casual\": 2, \"casual\": 1}\n"
    "_RAIN_CODES     = {51, 53, 55, 61, 63, 65, 80, 81, 82, 95, 96, 99}\n"
    "\n"
    "\n"
    "@router.get(\"/today\", response_model=AISuggestionResponse)\n"
    "def get_today_suggestion(\n"
    "    city:    Optional[str] = None,\n"
    "    db:      Session       = Depends(get_db),\n"
    "    user_id: Optional[int] = Depends(get_current_user_id),\n"
    "):\n"
    "    q = db.query(ClosetItemModel)\n"
    "    if user_id: q = q.filter(ClosetItemModel.user_id == user_id)\n"
    "    items = q.all()\n"
    "    if not items:\n"
    "        return AISuggestionResponse(\n"
    "            item_name=\"Empty closet\",\n"
    "            reason=\"Add items to your closet to get a daily suggestion.\")"
)
cb(
    "    today      = datetime.date.today().isoformat()\n"
    "    eq = db.query(CalendarEventModel).filter(\n"
    "             CalendarEventModel.start.like(f\"{today}%\"))\n"
    "    if user_id: eq = eq.filter(CalendarEventModel.user_id == user_id)\n"
    "    today_events = eq.all()\n"
    "    top_event    = (max(today_events,\n"
    "                        key=lambda e: _FORMALITY_RANK.get(e.formality or \"\", 0))\n"
    "                    if today_events else None)\n"
    "\n"
    "    weather  = _get_weather(city) if city else None\n"
    "    temp_c   = weather[\"temp_c\"]  if weather else None\n"
    "    is_rainy = weather and weather[\"code\"] in _RAIN_CODES\n"
    "    is_cold  = temp_c is not None and temp_c < 15\n"
    "    is_hot   = temp_c is not None and temp_c > 25\n"
    "\n"
    "    if   (top_event and top_event.formality == \"formal\"):\n"
    "        target = {\"FORMAL\", \"MODERATE\"}\n"
    "    elif (top_event and top_event.formality == \"business_casual\"):\n"
    "        target = {\"MODERATE\", \"FORMAL\", \"UNIVERSAL\"}\n"
    "    else:\n"
    "        target = {\"CASUAL\", \"MODERATE\", \"UNIVERSAL\"}\n"
    "\n"
    "    candidates = ([i for i in items\n"
    "                   if i.formality and i.formality.upper() in target]\n"
    "                  or items)\n"
    "\n"
    "    if   is_cold or is_rainy:\n"
    "        preferred = [i for i in candidates if i.category == \"Outerwear\"]\n"
    "    elif is_hot:\n"
    "        preferred = [i for i in candidates if i.category == \"Top\"]\n"
    "    else:\n"
    "        preferred = [i for i in candidates\n"
    "                     if i.category in (\"Top\", \"Outerwear\")]\n"
    "\n"
    "    pick = (preferred or candidates)[0]\n"
    "    reasons = []\n"
    "    if top_event: reasons.append(f\"you have '{top_event.title}' on your calendar\")\n"
    "    if temp_c is not None:\n"
    "        if is_rainy: reasons.append(f\"rain is expected ({temp_c:.0f}\\u00b0C)\")\n"
    "        elif is_cold: reasons.append(f\"it's cold at {temp_c:.0f}\\u00b0C\")\n"
    "        elif is_hot:  reasons.append(f\"it's warm at {temp_c:.0f}\\u00b0C\")\n"
    "        else:         reasons.append(f\"weather is mild at {temp_c:.0f}\\u00b0C\")\n"
    "    reason = (f\"Because {' and '.join(reasons)}, the {pick.name} is your best \"\n"
    "              f\"pick today.\" if reasons\n"
    "              else f\"The {pick.name} is a solid all-round choice for today.\")\n"
    "    return AISuggestionResponse(item_name=pick.name, reason=reason)"
)
cb(
    "@router.get(\"/gaps\")\n"
    "def get_wardrobe_gaps(\n"
    "    db:      Session       = Depends(get_db),\n"
    "    user_id: Optional[int] = Depends(get_current_user_id),\n"
    "):\n"
    "    q = db.query(ClosetItemModel)\n"
    "    if user_id: q = q.filter(ClosetItemModel.user_id == user_id)\n"
    "    items = q.all()\n"
    "\n"
    "    if len(items) < 5:\n"
    "        return {\"suggestion\": \"Keep building your closet!\",\n"
    "                \"reason\": \"Upload more items to get gap analysis.\"}\n"
    "\n"
    "    cats       = {\"Top\": 0, \"Bottom\": 0, \"Outerwear\": 0, \"Shoes\": 0}\n"
    "    formalities = {\"FORMAL\": 0, \"CASUAL\": 0, \"MODERATE\": 0}\n"
    "    for i in items:\n"
    "        if i.category  in cats:       cats[i.category]       += 1\n"
    "        if i.formality in formalities: formalities[i.formality] += 1\n"
    "\n"
    "    if cats[\"Bottom\"]   == 0:\n"
    "        return {\"suggestion\": \"Pants / Jeans\",\n"
    "                \"reason\": f\"You have {cats['Top']} tops but no bottoms! \"\n"
    "                           f\"Dark jeans or chinos unlock many outfits.\"}\n"
    "    if cats[\"Outerwear\"] == 0:\n"
    "        return {\"suggestion\": \"A versatile jacket\",\n"
    "                \"reason\": \"A denim jacket or blazer helps you layer \"\n"
    "                           \"for colder weather.\"}\n"
    "    if cats[\"Shoes\"]    == 0:\n"
    "        return {\"suggestion\": \"Everyday sneakers or loafers\",\n"
    "                \"reason\": \"Shoes are the foundation of your outfit.\"}\n"
    "    if formalities[\"FORMAL\"] == 0 and formalities[\"CASUAL\"] > 3:\n"
    "        return {\"suggestion\": \"A formal shirt or dress pants\",\n"
    "                \"reason\": \"Your wardrobe leans casual. One formal piece \"\n"
    "                           \"prepares you for unexpected events.\"}\n"
    "    ratio = cats[\"Top\"] / max(1, cats[\"Bottom\"])\n"
    "    if ratio > 4:\n"
    "        return {\"suggestion\": \"Versatile Bottoms\",\n"
    "                \"reason\": f\"You have a {ratio:.1f}x ratio of tops to bottoms. \"\n"
    "                           f\"One new pair creates {cats['Top']} new combinations.\"}\n"
    "    return {\"suggestion\": \"Statement Accessory\",\n"
    "            \"reason\": \"Your core wardrobe is balanced! \"\n"
    "                       \"A watch or belt adds personality.\"}"
)

# ── B.8 ──────────────────────────────────────────────────────────────────────
heading("B.8  Travel Packing Router: style_back/app/routers/travel.py", 2)
body(
    "The travel router provides a single endpoint, POST /travel/pack, which generates a "
    "personalised packing list and day-by-day outfit schedule for a specified destination "
    "and trip duration. The endpoint fetches the destination's current weather via the "
    "shared _get_weather() helper, uses the temperature to determine whether outerwear "
    "should be included, and then generates one outfit per travel day using a "
    "deterministic random seed derived from the destination name and trip length. This "
    "seeding strategy ensures that repeated calls for the same trip return identical "
    "results, which is important for user experience consistency. The packing list "
    "de-duplicates items across days using a set of item IDs, so a single pair of shoes "
    "that appears in every day's outfit is listed only once in the master packing list.",
    indent=0.5
)
cb(
    "import random\n"
    "from typing import Optional\n"
    "from fastapi import APIRouter, Depends, HTTPException\n"
    "from sqlalchemy.orm import Session\n"
    "from app.schemas import PackRequest, PackResponse, DailyOutfit, OutfitItem\n"
    "from app.database import get_db\n"
    "from app.models import ClosetItemModel\n"
    "from app.auth_utils import get_current_user_id\n"
    "from app.routers.weather import _get_weather\n"
    "\n"
    "router = APIRouter(prefix=\"/travel\", tags=[\"travel\"])\n"
    "\n"
    "\n"
    "@router.post(\"/pack\", response_model=PackResponse)\n"
    "def generate_packing_list(\n"
    "    req:     PackRequest,\n"
    "    db:      Session       = Depends(get_db),\n"
    "    user_id: Optional[int] = Depends(get_current_user_id),\n"
    "):\n"
    "    q = db.query(ClosetItemModel)\n"
    "    if user_id: q = q.filter(ClosetItemModel.user_id == user_id)\n"
    "    items = q.all()\n"
    "    if not items:\n"
    "        raise HTTPException(status_code=400,\n"
    "                            detail=\"Your closet is empty.\")\n"
    "\n"
    "    weather = _get_weather(req.destination)\n"
    "    temp_c  = weather[\"temp_c\"] if weather else None\n"
    "    if weather and temp_c is not None:\n"
    "        weather_summary = (f\"Expected weather in {weather['city']}: \"\n"
    "                           f\"{temp_c:.1f}\\u00b0C.\")\n"
    "        is_cold = temp_c < 18\n"
    "    else:\n"
    "        weather_summary = (f\"Could not fetch weather for \"\n"
    "                           f\"{req.destination}. Packing a balanced set.\")\n"
    "        is_cold = False"
)
cb(
    "    rng = random.Random(req.destination + str(req.days))\n"
    "\n"
    "    def pick(cat):\n"
    "        pool = [i for i in items if i.category == cat]\n"
    "        if not pool: return None\n"
    "        rng.shuffle(pool); return pool[0]\n"
    "\n"
    "    daily_outfits, packed_ids, packing_list = [], set(), []\n"
    "\n"
    "    for day in range(1, req.days + 1):\n"
    "        top    = pick(\"Top\")\n"
    "        bottom = pick(\"Bottom\")\n"
    "        shoes  = pick(\"Shoes\")\n"
    "        outer  = pick(\"Outerwear\") if is_cold else None\n"
    "\n"
    "        day_items = []\n"
    "        for i in [top, bottom, shoes, outer]:\n"
    "            if i:\n"
    "                oi = OutfitItem(id=i.id, name=i.name,\n"
    "                                category=i.category, image_url=i.image_url)\n"
    "                day_items.append(oi)\n"
    "                if i.id not in packed_ids:\n"
    "                    packed_ids.add(i.id)\n"
    "                    packing_list.append(oi)\n"
    "        daily_outfits.append(DailyOutfit(day=day, items=day_items))\n"
    "\n"
    "    return PackResponse(\n"
    "        destination=req.destination,\n"
    "        weather_summary=weather_summary,\n"
    "        packing_list=packing_list,\n"
    "        daily_outfits=daily_outfits,\n"
    "    )"
)

# ── B.9 ──────────────────────────────────────────────────────────────────────
pb()
heading("B.9  React Application Root: style_front/src/App.jsx", 2)
body(
    "App.jsx is the root component of the React 19 single-page application. It maintains "
    "the global authentication state and dark-mode preference in React state, both of "
    "which are initialised from localStorage on first render. The ProtectedRoute component "
    "wraps all authenticated views and redirects unauthenticated users to /login, "
    "preserving the originally requested URL in React Router's location state so the user "
    "is redirected back after login. Dark mode is implemented by toggling a 'dark' class "
    "on the document root element, which activates Tailwind CSS's dark: variant utilities "
    "throughout the application. The ClosetFilterProvider wraps all authenticated routes "
    "to provide a shared React context for the wardrobe filtering state, allowing the "
    "WardrobeView's filter bar and item grid to communicate without prop-drilling.",
    indent=0.5
)
cb(
    "import React, { useState, useEffect } from 'react';\n"
    "import { Routes, Route, Navigate,\n"
    "         useNavigate, useLocation } from 'react-router-dom';\n"
    "import LandingPage   from './pages/LandingPage';\n"
    "import AuthFlow      from './pages/AuthFlow';\n"
    "import DashboardView from './pages/DashboardView';\n"
    "import WardrobeView  from './pages/WardrobeView';\n"
    "import GeneratorView from './pages/GeneratorView';\n"
    "import TravelView    from './pages/TravelView';\n"
    "import ProfileView   from './pages/ProfileView';\n"
    "import AppShell      from './components/AppShell';\n"
    "import { ClosetFilterProvider } from './context/ClosetFilterContext';\n"
    "\n"
    "const ProtectedRoute = ({ children, isAuthenticated }) => {\n"
    "  const location = useLocation();\n"
    "  if (!isAuthenticated)\n"
    "    return <Navigate to=\"/login\" state={{ from: location }} replace />;\n"
    "  return children;\n"
    "};\n"
    "\n"
    "const DARK_KEY = 'style-dark-mode';\n"
    "\n"
    "const App = () => {\n"
    "  const [isAuthenticated, setIsAuthenticated] = useState(() => {\n"
    "    try { return Boolean(localStorage.getItem('auth_token')); }\n"
    "    catch { return false; }\n"
    "  });\n"
    "  const [activeTab, setActiveTab] = useState('wardrobe');\n"
    "  const [darkMode,  setDarkMode]  = useState(() => {\n"
    "    try { return localStorage.getItem(DARK_KEY) === '1'; }\n"
    "    catch { return false; }\n"
    "  });\n"
    "  const navigate = useNavigate();\n"
    "  const location = useLocation();"
)
cb(
    "  useEffect(() => {\n"
    "    document.documentElement.classList.toggle('dark', darkMode);\n"
    "    try { localStorage.setItem(DARK_KEY, darkMode ? '1' : '0'); }\n"
    "    catch (_) {}\n"
    "  }, [darkMode]);\n"
    "\n"
    "  // Sync sidebar highlight with current URL\n"
    "  useEffect(() => {\n"
    "    const p = location.pathname;\n"
    "    if      (p.includes('dashboard')) setActiveTab('dashboard');\n"
    "    else if (p.includes('wardrobe'))  setActiveTab('wardrobe');\n"
    "    else if (p.includes('generator')) setActiveTab('generator');\n"
    "    else if (p.includes('travel'))    setActiveTab('travel');\n"
    "    else if (p.includes('profile'))   setActiveTab('profile');\n"
    "  }, [location]);\n"
    "\n"
    "  const handleLogin = () => {\n"
    "    setIsAuthenticated(true);\n"
    "    navigate('/app/wardrobe');\n"
    "  };\n"
    "  const handleLogout = () => {\n"
    "    setIsAuthenticated(false);\n"
    "    try {\n"
    "      localStorage.removeItem('auth_token');\n"
    "      localStorage.removeItem('user_city');\n"
    "      localStorage.removeItem('user_name');\n"
    "    } catch {}\n"
    "    navigate('/');\n"
    "  };\n"
    "  const handleNavClick = (id) => {\n"
    "    setActiveTab(id);\n"
    "    navigate(`/app/${id}`);\n"
    "  };"
)
cb(
    "  return (\n"
    "    <Routes>\n"
    "      {/* Public routes */}\n"
    "      <Route path=\"/\"         element={<LandingPage />} />\n"
    "      <Route path=\"/login\"    element={<AuthFlow initialView=\"login\"\n"
    "                                         onComplete={handleLogin} />} />\n"
    "      <Route path=\"/register\" element={<AuthFlow initialView=\"signup\"\n"
    "                                         onComplete={handleLogin} />} />\n"
    "\n"
    "      {/* Protected app shell */}\n"
    "      <Route path=\"/app\" element={\n"
    "        <ProtectedRoute isAuthenticated={isAuthenticated}>\n"
    "          <ClosetFilterProvider>\n"
    "            <AppShell\n"
    "              activeTab={activeTab}\n"
    "              setActiveTab={handleNavClick}\n"
    "              onLogout={handleLogout}\n"
    "              darkMode={darkMode}\n"
    "              setDarkMode={setDarkMode}\n"
    "            />\n"
    "          </ClosetFilterProvider>\n"
    "        </ProtectedRoute>\n"
    "      }>\n"
    "        <Route path=\"dashboard\"  element={<DashboardView setView={handleNavClick} />} />\n"
    "        <Route path=\"wardrobe\"   element={<WardrobeView />} />\n"
    "        <Route path=\"generator\"  element={<GeneratorView />} />\n"
    "        <Route path=\"travel\"     element={<TravelView />} />\n"
    "        <Route path=\"profile\"    element={<ProfileView />} />\n"
    "        <Route index element={<Navigate to=\"wardrobe\" replace />} />\n"
    "      </Route>\n"
    "\n"
    "      <Route path=\"*\" element={<Navigate to=\"/\" replace />} />\n"
    "    </Routes>\n"
    "  );\n"
    "};\n"
    "\n"
    "export default App;"
)

# ── B.10 ─────────────────────────────────────────────────────────────────────
heading("B.10  Frontend API Client: style_front/src/api/client.js", 2)
body(
    "The client.js module is the single point of contact between the React frontend and "
    "the FastAPI backend. All HTTP communication is routed through the request() function, "
    "which reads the auth token from localStorage and attaches it as a Bearer token in "
    "the Authorization header. A lightweight in-memory cache (memCache object) is "
    "implemented via the cachedRequest() helper: on the first call, the response is "
    "fetched and stored; on subsequent calls, the cached value is returned immediately "
    "and a background re-fetch updates the cache without blocking the UI. This provides "
    "instant UI responses on page revisits while keeping data fresh. The clearCache() "
    "function is called by all mutation operations (create, update, delete) to ensure "
    "stale data is not served after a change. The BASE URL is injected at build time via "
    "the REACT_APP_API_URL environment variable, making the client environment-agnostic.",
    indent=0.5
)
cb(
    "/**\n"
    " * API client for Style backend (FastAPI).\n"
    " * Base URL: set REACT_APP_API_URL in .env\n"
    " */\n"
    "const BASE = process.env.REACT_APP_API_URL || '';\n"
    "\n"
    "function getAuthToken() {\n"
    "  if (typeof window === 'undefined') return null;\n"
    "  return window.localStorage.getItem('auth_token');\n"
    "}\n"
    "\n"
    "async function request(path, options = {}) {\n"
    "  const url   = `${BASE}${path}`;\n"
    "  const token = getAuthToken();\n"
    "  const res   = await fetch(url, {\n"
    "    headers: {\n"
    "      'Content-Type': 'application/json',\n"
    "      ...(token ? { Authorization: `Bearer ${token}` } : {}),\n"
    "      ...options.headers,\n"
    "    },\n"
    "    ...options,\n"
    "  });\n"
    "  if (!res.ok) {\n"
    "    const err = new Error(res.statusText);\n"
    "    err.status = res.status;\n"
    "    try   { err.body = await res.json(); }\n"
    "    catch { err.body = await res.text(); }\n"
    "    throw err;\n"
    "  }\n"
    "  return res.status === 204 ? null : res.json();\n"
    "}"
)
cb(
    "const memCache = {};\n"
    "\n"
    "async function cachedRequest(path, cacheKey) {\n"
    "  if (memCache[cacheKey]) {\n"
    "    // Background refresh to keep cache warm\n"
    "    request(path)\n"
    "      .then(res => { memCache[cacheKey] = res; })\n"
    "      .catch(() => {});\n"
    "    return Promise.resolve(memCache[cacheKey]);\n"
    "  }\n"
    "  const res = await request(path);\n"
    "  memCache[cacheKey] = res;\n"
    "  return res;\n"
    "}\n"
    "\n"
    "export function getCachedSync(cacheKey) {\n"
    "  return memCache[cacheKey] || null;\n"
    "}\n"
    "\n"
    "function clearCache() {\n"
    "  for (let key in memCache) delete memCache[key];\n"
    "}"
)
cb(
    "export const api = {\n"
    "  health() { return request('/health'); },\n"
    "\n"
    "  register({ name, email, password, city }) {\n"
    "    return request('/auth/register', {\n"
    "      method: 'POST',\n"
    "      body: JSON.stringify({ name, email, password, city }),\n"
    "    }).then(res => {\n"
    "      if (res?.access_token) {\n"
    "        localStorage.setItem('auth_token', res.access_token);\n"
    "        if (res.city) localStorage.setItem('user_city', res.city);\n"
    "        if (res.name) localStorage.setItem('user_name', res.name);\n"
    "      }\n"
    "      return res;\n"
    "    });\n"
    "  },\n"
    "\n"
    "  login({ email, password }) {\n"
    "    return request('/auth/login', {\n"
    "      method: 'POST',\n"
    "      body: JSON.stringify({ email, password }),\n"
    "    }).then(res => {\n"
    "      if (res?.access_token) {\n"
    "        localStorage.setItem('auth_token', res.access_token);\n"
    "        if (res.city) localStorage.setItem('user_city', res.city);\n"
    "        if (res.name) localStorage.setItem('user_name', res.name);\n"
    "      }\n"
    "      return res;\n"
    "    });\n"
    "  },\n"
    "\n"
    "  getProfile()          { return request('/auth/me'); },\n"
    "\n"
    "  updateProfile(data) {\n"
    "    return request('/auth/me', {\n"
    "      method: 'PUT', body: JSON.stringify(data),\n"
    "    }).then(res => {\n"
    "      if (res?.access_token) {\n"
    "        localStorage.setItem('auth_token', res.access_token);\n"
    "        localStorage.setItem('user_city',  res.city || '');\n"
    "        localStorage.setItem('user_name',  res.name || '');\n"
    "      }\n"
    "      return res;\n"
    "    });\n"
    "  },"
)
cb(
    "  getCloset({ category=null, color=null, formality=null, search=null } = {}) {\n"
    "    const params = new URLSearchParams();\n"
    "    if (category != null && category !== '' && category !== 'all')\n"
    "      params.set('category', category);\n"
    "    if (color    != null && color    !== '') params.set('color',    color);\n"
    "    if (formality!= null && formality!== '') params.set('formality',formality);\n"
    "    if (search   != null && search.trim()!== '') params.set('search', search.trim());\n"
    "    const q = params.toString() ? `?${params.toString()}` : '';\n"
    "    return cachedRequest(`/closet${q}`, `closet-${q}`);\n"
    "  },\n"
    "\n"
    "  createClosetItem(item) {\n"
    "    clearCache();\n"
    "    return request('/closet', { method: 'POST', body: JSON.stringify(item) });\n"
    "  },\n"
    "\n"
    "  uploadClosetItem(formData) {\n"
    "    clearCache();\n"
    "    const token = getAuthToken();\n"
    "    return fetch(`${BASE}/closet/upload`, {\n"
    "      method:  'POST',\n"
    "      headers: token ? { Authorization: `Bearer ${token}` } : {},\n"
    "      body:    formData,\n"
    "    }).then(async res => {\n"
    "      if (!res.ok) {\n"
    "        const err = new Error(res.statusText); err.status = res.status;\n"
    "        try { err.body = await res.json(); } catch { err.body = await res.text(); }\n"
    "        throw err;\n"
    "      }\n"
    "      return res.json();\n"
    "    });\n"
    "  },\n"
    "\n"
    "  updateClosetItem(id, data) {\n"
    "    clearCache();\n"
    "    return request(`/closet/${id}`, { method: 'PATCH', body: JSON.stringify(data) });\n"
    "  },\n"
    "\n"
    "  deleteClosetItem(id) {\n"
    "    clearCache();\n"
    "    return request(`/closet/${id}`, { method: 'DELETE' });\n"
    "  },"
)
cb(
    "  getTodayOutfit(refresh = false) {\n"
    "    return cachedRequest(\n"
    "      `/outfits/today${refresh ? '?refresh=true' : ''}`,\n"
    "      `outfit-today-${refresh}`);\n"
    "  },\n"
    "\n"
    "  generateOutfit({ context, weather_temp_c, formality_preference, vibe }) {\n"
    "    return request('/outfits/generate', {\n"
    "      method: 'POST',\n"
    "      body: JSON.stringify({\n"
    "        context:              context              || 'Office',\n"
    "        weather_temp_c:       weather_temp_c       ?? null,\n"
    "        formality_preference: formality_preference ?? null,\n"
    "        vibe:                 vibe                 ?? 50,\n"
    "      }),\n"
    "    });\n"
    "  },\n"
    "\n"
    "  saveOutfit({ context, items, explanation }) {\n"
    "    clearCache();\n"
    "    return request('/outfits/save', {\n"
    "      method: 'POST',\n"
    "      body: JSON.stringify({ context, items, explanation }),\n"
    "    });\n"
    "  },\n"
    "\n"
    "  getSavedOutfits()    { return cachedRequest('/outfits/saved', 'saved-outfits'); },\n"
    "\n"
    "  virtualTryOn(data) {\n"
    "    return request('/outfits/try-on',\n"
    "                   { method: 'POST', body: JSON.stringify(data) });\n"
    "  },\n"
    "\n"
    "  getWardrobeGaps() { return cachedRequest('/suggestions/gaps', 'wardrobe-gaps'); },\n"
    "\n"
    "  getCalendarEvents(date = null) {\n"
    "    const q = date ? `?date=${encodeURIComponent(date)}` : '';\n"
    "    return cachedRequest(`/calendar/events${q}`, `calendar-events-${q}`);\n"
    "  },\n"
    "\n"
    "  createCalendarEvent(event) {\n"
    "    clearCache();\n"
    "    return request('/calendar/events', { method: 'POST', body: JSON.stringify(event) });\n"
    "  },\n"
    "\n"
    "  getWeather({ city, lat, lon } = {}) {\n"
    "    const p = new URLSearchParams();\n"
    "    if (city) p.set('city', city);\n"
    "    if (lat  != null) p.set('lat', lat);\n"
    "    if (lon  != null) p.set('lon', lon);\n"
    "    return cachedRequest(`/weather/current?${p}`, `weather-${p}`);\n"
    "  },\n"
    "\n"
    "  generatePackingList(data) {\n"
    "    return request('/travel/pack',\n"
    "                   { method: 'POST', body: JSON.stringify(data) });\n"
    "  },\n"
    "};\n"
    "\n"
    "export const isApiConfigured = () => Boolean(BASE);"
)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX C — DATABASE SCHEMA
# ══════════════════════════════════════════════════════════════════════════════
pb()
heading("Appendix C: Database Schema Reference", 1)
body(
    "StylistAI uses SQLite as its development persistence layer, accessed exclusively "
    "through SQLAlchemy 2.0's ORM. The schema comprises four tables: users, closet_items, "
    "calendar_events, and saved_outfits. All tables are defined as SQLAlchemy ORM model "
    "classes in style_back/app/models.py and are auto-created at application startup via "
    "Base.metadata.create_all(). The following sections describe each table in detail, "
    "including column names, data types, constraints, and the rationale for the design "
    "decisions made. The schema is intentionally kept minimal to support the functional "
    "requirements of the current application version; a production deployment would "
    "expand it with additional indexes, foreign-key cascade rules, and audit columns.",
    indent=0.5
)

heading("C.1  Table: users", 2)
body(
    "The users table stores registered user accounts. Each row represents one user. "
    "The primary key is an auto-incremented integer id. The email column is subject to "
    "a UNIQUE constraint enforced at both the database level and the application level "
    "(checked in the register and update_me endpoints before committing). Passwords are "
    "never stored in plaintext; the password_hash column stores the bcrypt output, which "
    "includes the salt as a prefix and is safe to compare using bcrypt.checkpw(). The "
    "name and city columns are optional (nullable) and are used to personalise the "
    "dashboard greeting and to resolve the user's city for live weather lookups without "
    "requiring the client to send the city on every request.",
    indent=0.5
)
para()
t_users = doc.add_table(rows=1, cols=4)
t_users.style = "Table Grid"
for cell, txt in zip(t_users.rows[0].cells, ["Column", "Type", "Constraints", "Description"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(10)
for row in [
    ("id",            "INTEGER",  "PRIMARY KEY, AUTOINCREMENT", "Surrogate primary key"),
    ("email",         "VARCHAR",  "NOT NULL, UNIQUE",           "User's login email address"),
    ("password_hash", "VARCHAR",  "NOT NULL",                   "bcrypt hash (salt + hash concatenated)"),
    ("name",          "VARCHAR",  "NULLABLE",                   "Display name shown in dashboard greeting"),
    ("city",          "VARCHAR",  "NULLABLE",                   "Home city for live weather lookups"),
]:
    add_table_row(t_users, row)
para()

heading("C.2  Table: closet_items", 2)
body(
    "The closet_items table stores individual garment records. The user_id column is a "
    "foreign key referencing users.id, establishing the ownership relationship that "
    "allows each user to see only their own items. The category column stores one of the "
    "five enum values defined in the Category schema (Top, Bottom, Outerwear, Shoes, "
    "Accessory) as a plain string, avoiding the need for a separate categories lookup "
    "table. The image_url column stores the fully-qualified URL of the uploaded and "
    "background-removed garment image, served by FastAPI's StaticFiles middleware. The "
    "formality column stores one of four string values (CASUAL, MODERATE, FORMAL, "
    "UNIVERSAL) that are used by the outfit recommendation engine to filter items by "
    "appropriateness. The formality_value column stores a 0-100 integer for progress "
    "bar display in the frontend.",
    indent=0.5
)
para()
t_closet = doc.add_table(rows=1, cols=4)
t_closet.style = "Table Grid"
for cell, txt in zip(t_closet.rows[0].cells, ["Column", "Type", "Constraints", "Description"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(10)
for row in [
    ("id",             "INTEGER",  "PRIMARY KEY, AUTOINCREMENT", "Surrogate primary key"),
    ("user_id",        "INTEGER",  "FK → users.id, NULLABLE",    "Owner; NULL allows anonymous items"),
    ("name",           "VARCHAR",  "NOT NULL",                   "User-given garment name"),
    ("category",       "VARCHAR",  "NOT NULL",                   "Top | Bottom | Outerwear | Shoes | Accessory"),
    ("image_url",      "VARCHAR",  "NULLABLE",                   "Full URL to processed PNG image"),
    ("color",          "VARCHAR",  "NULLABLE",                   "Primary colour keyword (black, navy, etc.)"),
    ("formality",      "VARCHAR",  "NULLABLE",                   "CASUAL | MODERATE | FORMAL | UNIVERSAL"),
    ("formality_value","INTEGER",  "NULLABLE",                   "0-100 numeric formality score"),
]:
    add_table_row(t_closet, row)
para()

heading("C.3  Table: calendar_events", 2)
body(
    "The calendar_events table stores user-defined schedule events that inform the outfit "
    "recommendation engine. The start and end columns store ISO 8601 datetime strings "
    "(e.g. '2026-05-07T09:00') rather than native DATETIME values, simplifying "
    "serialisation and comparison. The formality column accepts three values — formal, "
    "business_casual, and casual — and is used to determine the formality target set "
    "during outfit generation. The recommendation engine selects the highest-priority "
    "event of the day using the _FORMALITY_RANK dictionary, so a formal meeting "
    "overrides a casual coffee catch-up when both are on the same day.",
    indent=0.5
)
para()
t_cal = doc.add_table(rows=1, cols=4)
t_cal.style = "Table Grid"
for cell, txt in zip(t_cal.rows[0].cells, ["Column", "Type", "Constraints", "Description"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(10)
for row in [
    ("id",       "VARCHAR",  "PRIMARY KEY",          "Client-generated UUID string"),
    ("user_id",  "INTEGER",  "FK → users.id, NULLABLE","Owner of the calendar event"),
    ("title",    "VARCHAR",  "NOT NULL",              "Event display name"),
    ("start",    "VARCHAR",  "NOT NULL",              "ISO 8601 start datetime string"),
    ("end",      "VARCHAR",  "NOT NULL",              "ISO 8601 end datetime string"),
    ("formality","VARCHAR",  "NULLABLE",              "formal | business_casual | casual"),
]:
    add_table_row(t_cal, row)
para()

heading("C.4  Table: saved_outfits", 2)
body(
    "The saved_outfits table persists outfits that the user has explicitly chosen to save. "
    "Because an outfit is a variable-length collection of closet items, the items are "
    "serialised as a JSON array and stored in the items_json TEXT column rather than in "
    "a separate junction table. Each JSON element is a snapshot of the OutfitItem schema "
    "at the time of saving, meaning that saved outfits remain intact even if the "
    "underlying closet items are later deleted or modified. The saved_at column is "
    "populated automatically by a SQLAlchemy server_default, recording the UTC timestamp "
    "at which the outfit was saved.",
    indent=0.5
)
para()
t_saved = doc.add_table(rows=1, cols=4)
t_saved.style = "Table Grid"
for cell, txt in zip(t_saved.rows[0].cells, ["Column", "Type", "Constraints", "Description"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(10)
for row in [
    ("id",          "VARCHAR",  "PRIMARY KEY",              "UUID string generated by the server"),
    ("user_id",     "INTEGER",  "FK → users.id, NULLABLE",  "Owner of the saved outfit"),
    ("context",     "VARCHAR",  "NOT NULL",                 "Occasion label (Office, Gym, etc.)"),
    ("items_json",  "TEXT",     "NOT NULL",                 "JSON array of OutfitItem objects (snapshot)"),
    ("explanation", "TEXT",     "NOT NULL",                 "Natural-language explanation of the outfit"),
    ("saved_at",    "DATETIME", "server_default=NOW()",     "UTC timestamp when outfit was saved"),
]:
    add_table_row(t_saved, row)
para()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX D — FRONTEND COMPONENT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
pb()
heading("Appendix D: Frontend Component Architecture", 1)
body(
    "The StylistAI frontend is a React 19 single-page application built with Vite and "
    "styled with Tailwind CSS v3. The component tree is organised into four layers: "
    "routing (App.jsx), shell (AppShell.jsx), page views (five protected view components), "
    "and shared UI primitives (Toast, modals, cards). This section describes each "
    "significant component, its responsibilities, and the key state it manages.",
    indent=0.5
)

heading("D.1  AppShell Component", 2)
body(
    "AppShell is the authenticated application frame. It renders the persistent left "
    "sidebar navigation on desktop (which collapses to a bottom tab bar on mobile), the "
    "top header bar with the dark-mode toggle and logout button, and the <Outlet /> "
    "placeholder from React Router where the active page view is rendered. AppShell "
    "receives the activeTab, setActiveTab, onLogout, darkMode, and setDarkMode props "
    "from App.jsx and passes them down to the sidebar and header sub-components. The "
    "sidebar highlights the active tab by comparing the tab id with the activeTab prop "
    "and applying a conditional Tailwind class.",
    indent=0.5
)

heading("D.2  DashboardView Component", 2)
body(
    "DashboardView is the home screen shown immediately after login. It orchestrates "
    "four parallel data fetches on mount: the today's outfit from GET /outfits/today, "
    "the current weather from GET /weather/current using the user's stored city, the "
    "day's calendar events from GET /calendar/events, and the wardrobe gap analysis from "
    "GET /suggestions/gaps. These are fetched independently using separate useEffect and "
    "api.* calls, and their results are held in separate useState variables. This "
    "parallel loading approach ensures the page is interactive as soon as each "
    "individual data source resolves, rather than waiting for all fetches to complete. "
    "The outfit card renders a 2x2 image grid of garment thumbnails when four items are "
    "available, collapsing to a single column for fewer items. The gap analysis card is "
    "conditionally rendered only when the closet contains items, since gap analysis on "
    "an empty wardrobe is not meaningful.",
    indent=0.5
)

heading("D.3  WardrobeView Component", 2)
body(
    "WardrobeView displays the user's complete digital closet as a responsive grid of "
    "garment cards. It reads filter state from the ClosetFilterContext (category, colour, "
    "formality, and free-text search) and passes the active filters to api.getCloset() "
    "on every filter change. Each garment card displays the background-removed PNG image, "
    "the item name, category badge, formality pill, and a colour swatch. Clicking a card "
    "opens an edit modal where the user can update the item's name, category, colour, "
    "and formality, or delete the item. The upload panel supports both file selection and "
    "drag-and-drop, submitting the image and metadata as a multipart form to "
    "POST /closet/upload. A loading skeleton grid is displayed while items are being "
    "fetched, improving perceived performance.",
    indent=0.5
)

heading("D.4  GeneratorView Component", 2)
body(
    "GeneratorView provides the manual outfit generation interface. The component guides "
    "the user through a two-step flow: first selecting an occasion from a grid of six "
    "preset tiles (Office, Date Night, Travel, Gym, Casual, Special Event), then "
    "adjusting a vibe slider from 0 (Pure Comfort) to 100 (Style Focus). Generating an "
    "outfit calls POST /outfits/generate with the selected context and vibe score, plus "
    "the current weather temperature if it was successfully fetched from the weather "
    "endpoint. The result is displayed as a list of item cards with formality badges. "
    "From the results screen the user can save the outfit (POST /outfits/save) or trigger "
    "the DALL-E 3 virtual try-on (POST /outfits/try-on). The try-on generates a "
    "1024x1024 portrait image and displays it in a modal overlay. Saved outfit history "
    "is displayed below the generator in a scrollable carousel.",
    indent=0.5
)

heading("D.5  TravelView Component", 2)
body(
    "TravelView is the Smart Packing List interface. The user enters a destination city "
    "and trip duration in days, then submits the form to POST /travel/pack. The response "
    "includes a weather summary for the destination, a deduplicated master packing list, "
    "and a day-by-day outfit breakdown. These are rendered as an expandable accordion: "
    "the packing list is shown at the top as a checklist that users can tick off, "
    "followed by one collapsible section per travel day. Each day section shows the four "
    "outfit items as small image thumbnails with their names and categories. The "
    "destination weather summary is shown as a prominent banner at the top of the results "
    "to help the user understand why particular items were selected.",
    indent=0.5
)

heading("D.6  ProfileView Component", 2)
body(
    "ProfileView allows the authenticated user to update their display name, home city, "
    "email address, and password. The form pre-populates with the values retrieved from "
    "GET /auth/me on mount. On submission, the changed fields are sent to PUT /auth/me. "
    "If the server returns a new access token (which occurs when the email is changed, "
    "since the token encodes the email), the new token is written to localStorage and "
    "the auth context is updated so subsequent API calls use the updated credentials. "
    "The city field is particularly important as it determines which city's weather is "
    "used for the dashboard outfit recommendation.",
    indent=0.5
)

heading("D.7  Context: ClosetFilterContext", 2)
body(
    "ClosetFilterContext is a React context object that provides the wardrobe filter state "
    "(category, colour, formality, search query) and its setter functions to any "
    "descendant component in the authenticated route tree. This avoids the need to prop-"
    "drill filter state through AppShell down to WardrobeView. The context also holds a "
    "refresh counter that is incremented whenever a mutation occurs (upload, update, "
    "delete), causing WardrobeView to re-fetch the closet list and display the updated "
    "items without requiring a full page reload.",
    indent=0.5
)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX E — ENVIRONMENT SETUP AND DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
pb()
heading("Appendix E: Environment Setup and Deployment Guide", 1)
body(
    "This appendix provides a complete step-by-step guide to setting up the StylistAI "
    "development environment from a clean machine. Both the backend and frontend must be "
    "running simultaneously for the application to function. The backend serves the "
    "REST API and static uploaded files; the frontend serves the React SPA from the Vite "
    "development server.",
    indent=0.5
)

heading("E.1  Prerequisites", 2)
bullet("Python 3.11 or later (tested with Python 3.12)")
bullet("Node.js 20 or later with npm 10 or later")
bullet("An OpenAI API key with access to the DALL-E 3 image generation model")
bullet("Approximately 1 GB of free disk space (the rembg U2-Net model is ~170 MB; "
       "npm dependencies are ~500 MB)")

heading("E.2  Backend Setup", 2)
body("Navigate to the style_back directory and create a virtual environment:", indent=0.5)
cb(
    "cd style_back\n"
    "python3 -m venv .venv\n"
    "source .venv/bin/activate          # Windows: .venv\\Scripts\\activate\n"
    "pip install -r requirements.txt"
)
body("Create a .env file in style_back/ with the following contents:", indent=0.5)
cb(
    "DATABASE_URL=sqlite:///./style.db\n"
    "OPENAI_API_KEY=sk-..."
)
body("Start the backend development server:", indent=0.5)
cb("uvicorn main:app --reload --port 8000")
body(
    "The API will be available at http://localhost:8000 and the interactive Swagger "
    "documentation at http://localhost:8000/docs. On first startup, SQLAlchemy will "
    "create the style.db SQLite database file and all four tables automatically. The "
    "rembg U2-Net model will be downloaded (~170 MB) on the first image upload request.",
    indent=0.5
)

heading("E.3  Frontend Setup", 2)
body("Navigate to the style_front directory and install dependencies:", indent=0.5)
cb(
    "cd style_front\n"
    "npm install"
)
body("Create a .env file in style_front/ with the following contents:", indent=0.5)
cb("REACT_APP_API_URL=http://localhost:8000")
body("Start the frontend development server:", indent=0.5)
cb("npm start")
body(
    "The React application will be available at http://localhost:3000. The Vite dev "
    "server proxies API requests to the backend and supports hot module replacement "
    "for rapid development iteration.",
    indent=0.5
)

heading("E.4  Project Directory Structure", 2)
body(
    "The following listing shows the complete directory structure of the StylistAI "
    "project as submitted:",
    indent=0.5
)
cb(
    "StylistAI/\n"
    "├── style_back/                 # FastAPI backend\n"
    "│   ├── main.py                 # ASGI app entry point\n"
    "│   ├── requirements.txt        # Python dependencies\n"
    "│   ├── .env                    # Environment variables (not committed)\n"
    "│   ├── style.db                # SQLite database (auto-created)\n"
    "│   ├── uploads/                # Uploaded garment images (auto-created)\n"
    "│   └── app/\n"
    "│       ├── database.py         # SQLAlchemy engine & session\n"
    "│       ├── models.py           # ORM model classes\n"
    "│       ├── schemas.py          # Pydantic v2 schemas\n"
    "│       ├── auth_utils.py       # Token → user resolution\n"
    "│       └── routers/\n"
    "│           ├── auth.py         # /auth endpoints\n"
    "│           ├── closet.py       # /closet CRUD + upload\n"
    "│           ├── outfits.py      # /outfits generation & try-on\n"
    "│           ├── suggestions.py  # /suggestions today & gaps\n"
    "│           ├── travel.py       # /travel/pack\n"
    "│           ├── calendar.py     # /calendar/events CRUD\n"
    "│           ├── weather.py      # /weather/current\n"
    "│           └── health.py       # /health\n"
    "└── style_front/                # React frontend\n"
    "    ├── package.json\n"
    "    ├── .env\n"
    "    ├── tailwind.config.js\n"
    "    └── src/\n"
    "        ├── App.jsx             # Router root\n"
    "        ├── api/\n"
    "        │   └── client.js       # HTTP client + cache\n"
    "        ├── components/\n"
    "        │   ├── AppShell.jsx    # Sidebar + header frame\n"
    "        │   └── Toast.jsx       # Toast notification\n"
    "        ├── context/\n"
    "        │   └── ClosetFilterContext.jsx\n"
    "        └── pages/\n"
    "            ├── LandingPage.jsx\n"
    "            ├── AuthFlow.jsx\n"
    "            ├── DashboardView.jsx\n"
    "            ├── WardrobeView.jsx\n"
    "            ├── GeneratorView.jsx\n"
    "            ├── TravelView.jsx\n"
    "            └── ProfileView.jsx"
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(DOCX)
print(f"Saved expanded report: {DOCX}")

# Quick word count
from docx import Document as D2
d = D2(DOCX)
words = sum(len(p.text.split()) for p in d.paragraphs if p.text.strip())
print(f"Estimated word count: {words}")
print(f"Approx pages at 350 words/page: ~{words // 350}")
