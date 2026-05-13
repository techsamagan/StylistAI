"""
Second expansion pass: adds Appendix F (formal SRS), Appendix G (test suite code),
and Appendix H (UI walkthrough) to push the report past 60 pages.
Run AFTER expand_report.py.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = "/Users/samagannurdinov/Desktop/StylistAI/final/StylistAI_Final_Report_Samagan_Nurdinov.docx"
doc  = Document(DOCX)

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, before=24, after=8, size=None, color=None):
    sz = (size or 16) if level == 1 else (size or 14) if level == 2 else (size or 12)
    p  = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    for run in p.runs:
        run.font.name      = "Times New Roman"
        run.font.size      = Pt(sz)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(*(color or (0, 0, 0)))
    return p

def body(text, before=0, after=6, indent=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = Pt(22)
    if indent:
        p.paragraph_format.first_line_indent = Inches(indent)
    r = p.add_run(text)
    set_font(r, size=12, bold=bold)
    return p

def bullet(text, before=2, after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    set_font(r, size=12)
    return p

def cb(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    return p

def pb():
    doc.add_page_break()

def para(text="", before=0, after=0, size=12, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic)
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
    return row

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX F — SOFTWARE REQUIREMENTS SPECIFICATION
# ══════════════════════════════════════════════════════════════════════════════
pb()
heading("Appendix F: Software Requirements Specification (SRS)", 1)
body(
    "This appendix presents the formal Software Requirements Specification for StylistAI, "
    "written in accordance with the IEEE 830-1998 standard for software requirements "
    "documentation. The SRS defines the complete set of functional and non-functional "
    "requirements that were used to guide development and validate the final system. Each "
    "requirement is assigned a unique identifier (FR-XXX for functional, NFR-XXX for "
    "non-functional) and a priority level (Must Have, Should Have, Could Have) following "
    "the MoSCoW prioritisation framework. The requirements were elicited through a "
    "combination of literature review (identifying features common in state-of-the-art "
    "wardrobe applications), domain analysis of user needs in personal styling, and "
    "iterative refinement during the development process.",
    indent=0.5
)

heading("F.1  Functional Requirements", 2)
body(
    "The following functional requirements define the observable behaviours the system "
    "must exhibit. They are grouped by module and ordered by the sequence in which they "
    "are typically exercised by a new user.",
    indent=0.5
)

body("F.1.1  Authentication and User Management", bold=True, before=8)
para()
t_fr1 = doc.add_table(rows=1, cols=4)
t_fr1.style = "Table Grid"
for cell, txt in zip(t_fr1.rows[0].cells, ["ID", "Requirement", "Priority", "Acceptance Criterion"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)
for row in [
    ("FR-001", "The system shall allow a new user to register with email, password, optional name, and optional city.", "Must Have", "POST /auth/register returns 201 with access token; duplicate email returns 400."),
    ("FR-002", "The system shall authenticate registered users using email and password credentials.", "Must Have", "POST /auth/login returns 200 with access token for valid credentials; returns 400 for invalid."),
    ("FR-003", "The system shall store passwords using bcrypt hashing with a randomly generated salt.", "Must Have", "Database password_hash column contains bcrypt-formatted string; plaintext never stored."),
    ("FR-004", "The system shall allow authenticated users to retrieve their profile (name, email, city).", "Must Have", "GET /auth/me returns 200 with UserProfile for valid token."),
    ("FR-005", "The system shall allow authenticated users to update their name, city, email, or password.", "Must Have", "PUT /auth/me with partial fields updates only the provided fields."),
    ("FR-006", "The system shall reject profile updates that would assign an already-registered email to another account.", "Must Have", "PUT /auth/me with existing email returns 400."),
]:
    r = t_fr1.add_row()
    for i, val in enumerate(row):
        r.cells[i].text = val
        for p in r.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(9)
para()

body("F.1.2  Closet Management Module", bold=True, before=8)
para()
t_fr2 = doc.add_table(rows=1, cols=4)
t_fr2.style = "Table Grid"
for cell, txt in zip(t_fr2.rows[0].cells, ["ID", "Requirement", "Priority", "Acceptance Criterion"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)
for row in [
    ("FR-010", "The system shall allow users to add garment items with name, category, color, and formality metadata.", "Must Have", "POST /closet returns 201 with ClosetItem; item appears in subsequent GET /closet."),
    ("FR-011", "The system shall accept garment image uploads and store them as static files accessible via URL.", "Must Have", "POST /closet/upload with file returns 201; image_url is a valid reachable URL."),
    ("FR-012", "The system shall apply AI background removal to uploaded garment images using the U2-Net model.", "Should Have", "Uploaded PNG images have transparent backgrounds in ≥90% of test cases."),
    ("FR-013", "The system shall list all closet items for the authenticated user, supporting optional filters.", "Must Have", "GET /closet with no filters returns all user items; category filter returns only matching items."),
    ("FR-014", "The system shall allow partial updates to closet item metadata using HTTP PATCH.", "Must Have", "PATCH /closet/{id} with one field updates only that field; others remain unchanged."),
    ("FR-015", "The system shall allow users to delete individual closet items.", "Must Have", "DELETE /closet/{id} returns 204; item absent from subsequent GET /closet."),
    ("FR-016", "The system shall isolate each user's closet items so that no user can view or modify another user's items.", "Must Have", "Requests authenticated as User A return only User A's items."),
]:
    r = t_fr2.add_row()
    for i, val in enumerate(row):
        r.cells[i].text = val
        for p in r.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(9)
para()

body("F.1.3  Outfit Recommendation Module", bold=True, before=8)
para()
t_fr3 = doc.add_table(rows=1, cols=4)
t_fr3.style = "Table Grid"
for cell, txt in zip(t_fr3.rows[0].cells, ["ID", "Requirement", "Priority", "Acceptance Criterion"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)
for row in [
    ("FR-020", "The system shall generate an outfit suggestion given a context label and vibe score (0-100).", "Must Have", "POST /outfits/generate returns OutfitSuggestion with ≥1 item and non-empty explanation."),
    ("FR-021", "The vibe score shall influence formality: scores <30 select casual items; scores >70 select formal items.", "Must Have", "Vibe=10 outfit contains only CASUAL/MODERATE items; vibe=90 contains FORMAL/MODERATE items."),
    ("FR-022", "The system shall include outerwear in generated outfits when temperature is below 18°C.", "Should Have", "Outfit generated with weather_temp_c=10 includes an Outerwear item if available."),
    ("FR-023", "The system shall generate a daily outfit using live weather data and the user's calendar events.", "Must Have", "GET /outfits/today returns OutfitSuggestion; explanation references weather city when city is configured."),
    ("FR-024", "The daily outfit shall remain consistent across multiple calls on the same day (deterministic).", "Must Have", "Two GET /outfits/today calls without ?refresh return identical item sets."),
    ("FR-025", "The system shall return a different outfit when ?refresh=true is passed.", "Should Have", "GET /outfits/today?refresh=true has a reasonable chance of returning a different set."),
    ("FR-026", "The system shall allow users to save generated outfits for future reference.", "Must Have", "POST /outfits/save returns SavedOutfit; outfit appears in GET /outfits/saved."),
    ("FR-027", "The system shall allow users to delete saved outfits.", "Must Have", "DELETE /outfits/saved/{id} returns 204; outfit absent from subsequent GET /outfits/saved."),
    ("FR-028", "The system shall generate a photorealistic virtual try-on image using DALL-E 3.", "Should Have", "POST /outfits/try-on returns a non-empty image_url when OPENAI_API_KEY is configured."),
]:
    r = t_fr3.add_row()
    for i, val in enumerate(row):
        r.cells[i].text = val
        for p in r.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(9)
para()

body("F.1.4  Travel Packing Module", bold=True, before=8)
para()
t_fr4 = doc.add_table(rows=1, cols=4)
t_fr4.style = "Table Grid"
for cell, txt in zip(t_fr4.rows[0].cells, ["ID", "Requirement", "Priority", "Acceptance Criterion"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)
for row in [
    ("FR-030", "The system shall generate a day-by-day packing plan for a given destination and trip duration.", "Must Have", "POST /travel/pack with destination and days returns PackResponse with daily_outfits of length=days."),
    ("FR-031", "The packing plan shall include a weather summary for the destination fetched from Open-Meteo.", "Should Have", "weather_summary in response contains city name and temperature when destination is valid."),
    ("FR-032", "The packing list shall include outerwear when the destination temperature is below 18°C.", "Should Have", "Pack for a cold destination (e.g. Reykjavik in winter) includes Outerwear items."),
    ("FR-033", "The packing list shall deduplicate items that appear across multiple days.", "Must Have", "A 5-day trip packing list contains each item at most once."),
    ("FR-034", "The system shall return 400 if the user requests a packing list with an empty closet.", "Must Have", "POST /travel/pack with empty closet returns HTTP 400 with descriptive detail."),
]:
    r = t_fr4.add_row()
    for i, val in enumerate(row):
        r.cells[i].text = val
        for p in r.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(9)
para()

heading("F.2  Non-Functional Requirements", 2)
para()
t_nfr = doc.add_table(rows=1, cols=4)
t_nfr.style = "Table Grid"
for cell, txt in zip(t_nfr.rows[0].cells, ["ID", "Requirement", "Category", "Acceptance Criterion"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)
for row in [
    ("NFR-001", "Non-try-on API endpoints shall respond within 500 ms under single-user load.", "Performance", "Measured via curl on a standard development machine; average <500 ms across 20 calls."),
    ("NFR-002", "The React frontend shall achieve a Lighthouse Performance score ≥70 on the initial page load.", "Performance", "Chrome Lighthouse audit on the landing page."),
    ("NFR-003", "Passwords shall never be logged, returned in API responses, or stored in plaintext.", "Security", "Database inspection, API response inspection, and log inspection show no plaintext passwords."),
    ("NFR-004", "The application shall function correctly on Chrome 120+, Firefox 121+, and Safari 17+.", "Compatibility", "Manual smoke test on each browser passes all five views."),
    ("NFR-005", "The frontend shall be fully responsive at viewport widths from 375 px (iPhone SE) to 1920 px.", "Usability", "All five views render without horizontal scroll at 375 px and 1920 px."),
    ("NFR-006", "The application shall support dark mode toggling that persists across sessions.", "Usability", "Dark mode state read from and written to localStorage; toggle applies immediately."),
    ("NFR-007", "All API endpoints shall be documented in the auto-generated OpenAPI schema at /docs.", "Maintainability", "FastAPI Swagger UI at /docs lists all 25 endpoints with correct request/response schemas."),
    ("NFR-008", "The backend shall be database-agnostic: switching from SQLite to PostgreSQL shall require only a DATABASE_URL environment variable change.", "Portability", "Code review confirms no SQLite-specific SQL; engine branch in database.py supports both."),
    ("NFR-009", "The system shall handle Open-Meteo API failures gracefully without crashing.", "Reliability", "With network access blocked, GET /outfits/today returns a valid outfit without weather context."),
    ("NFR-010", "The rembg background removal shall degrade gracefully if the U2-Net model fails to load.", "Reliability", "Uploading an image with rembg unavailable saves the raw image without returning an error."),
]:
    r = t_nfr.add_row()
    for i, val in enumerate(row):
        r.cells[i].text = val
        for p in r.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(9)
para()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX G — TEST SUITE CODE
# ══════════════════════════════════════════════════════════════════════════════
pb()
heading("Appendix G: Test Suite Source Code", 1)
body(
    "This appendix presents the complete source code of the StylistAI pytest test suite. "
    "Tests are located in style_back/tests/ and are executed with the command "
    "'pytest tests/ -v' from the style_back directory. The test suite uses FastAPI's "
    "TestClient for HTTP-layer integration tests and an in-memory SQLite database "
    "(configured via the override_get_db fixture) to ensure complete test isolation: "
    "each test function receives a fresh, empty database with no shared state. The "
    "conftest.py module defines the shared fixtures used across all test modules.",
    indent=0.5
)

heading("G.1  Test Configuration: tests/conftest.py", 2)
body(
    "The conftest.py file defines the database fixture that creates a fresh in-memory "
    "SQLite database for every test run, the client fixture that patches the app's "
    "get_db dependency to use the test database, and the auth_headers fixture that "
    "registers a test user and returns an Authorization header for authenticated "
    "endpoints. Using an in-memory database (sqlite:///:memory:) ensures tests are "
    "completely isolated from the development database and run without disk I/O.",
    indent=0.5
)
cb(
    "# tests/conftest.py\n"
    "import pytest\n"
    "from fastapi.testclient import TestClient\n"
    "from sqlalchemy import create_engine\n"
    "from sqlalchemy.orm import sessionmaker\n"
    "from app.database import Base, get_db\n"
    "from main import app\n"
    "\n"
    "TEST_DB_URL = \"sqlite:///:memory:\"\n"
    "\n"
    "@pytest.fixture(scope=\"function\")\n"
    "def db_engine():\n"
    "    engine = create_engine(\n"
    "        TEST_DB_URL,\n"
    "        connect_args={\"check_same_thread\": False}\n"
    "    )\n"
    "    Base.metadata.create_all(bind=engine)\n"
    "    yield engine\n"
    "    Base.metadata.drop_all(bind=engine)\n"
    "    engine.dispose()\n"
    "\n"
    "\n"
    "@pytest.fixture(scope=\"function\")\n"
    "def client(db_engine):\n"
    "    TestSession = sessionmaker(\n"
    "        autocommit=False, autoflush=False, bind=db_engine)\n"
    "\n"
    "    def override_get_db():\n"
    "        db = TestSession()\n"
    "        try:\n"
    "            yield db\n"
    "        finally:\n"
    "            db.close()\n"
    "\n"
    "    app.dependency_overrides[get_db] = override_get_db\n"
    "    with TestClient(app) as c:\n"
    "        yield c\n"
    "    app.dependency_overrides.clear()\n"
    "\n"
    "\n"
    "@pytest.fixture\n"
    "def auth_headers(client):\n"
    "    client.post(\"/auth/register\", json={\n"
    "        \"email\":    \"test@example.com\",\n"
    "        \"password\": \"securepass123\",\n"
    "        \"name\":     \"Test User\",\n"
    "        \"city\":     \"London\",\n"
    "    })\n"
    "    return {\"Authorization\": \"Bearer token-test@example.com\"}\n"
    "\n"
    "\n"
    "@pytest.fixture\n"
    "def sample_item(client, auth_headers):\n"
    "    res = client.post(\"/closet\", json={\n"
    "        \"name\":     \"White Oxford Shirt\",\n"
    "        \"category\": \"Top\",\n"
    "        \"color\":    \"white\",\n"
    "        \"formality\": \"FORMAL\",\n"
    "        \"formality_value\": 80,\n"
    "    }, headers=auth_headers)\n"
    "    return res.json()"
)

heading("G.2  Authentication Tests: tests/test_auth.py", 2)
body(
    "The authentication test module verifies the register, login, and profile endpoints. "
    "The tests cover both the happy-path flows and the key error cases: duplicate email "
    "registration, incorrect password login, and unauthenticated profile access. Each "
    "test function uses the client fixture and is therefore fully isolated.",
    indent=0.5
)
cb(
    "# tests/test_auth.py\n"
    "import pytest\n"
    "\n"
    "\n"
    "class TestRegister:\n"
    "    def test_register_returns_201_with_token(self, client):\n"
    "        res = client.post(\"/auth/register\", json={\n"
    "            \"email\": \"alice@example.com\",\n"
    "            \"password\": \"pass123\",\n"
    "        })\n"
    "        assert res.status_code == 201\n"
    "        data = res.json()\n"
    "        assert \"access_token\" in data\n"
    "        assert data[\"token_type\"] == \"bearer\"\n"
    "\n"
    "    def test_register_stores_name_and_city(self, client):\n"
    "        res = client.post(\"/auth/register\", json={\n"
    "            \"email\": \"bob@example.com\",\n"
    "            \"password\": \"pass123\",\n"
    "            \"name\":     \"Bob Smith\",\n"
    "            \"city\":     \"New York\",\n"
    "        })\n"
    "        assert res.status_code == 201\n"
    "        data = res.json()\n"
    "        assert data[\"name\"] == \"Bob Smith\"\n"
    "        assert data[\"city\"] == \"New York\"\n"
    "\n"
    "    def test_register_duplicate_email_returns_400(self, client):\n"
    "        payload = {\"email\": \"dup@example.com\", \"password\": \"pass\"}\n"
    "        client.post(\"/auth/register\", json=payload)\n"
    "        res = client.post(\"/auth/register\", json=payload)\n"
    "        assert res.status_code == 400\n"
    "        assert \"already exists\" in res.json()[\"detail\"].lower()\n"
    "\n"
    "\n"
    "class TestLogin:\n"
    "    def test_login_valid_credentials_returns_token(self, client):\n"
    "        client.post(\"/auth/register\", json={\n"
    "            \"email\": \"carol@example.com\", \"password\": \"mypassword\"})\n"
    "        res = client.post(\"/auth/login\", json={\n"
    "            \"email\": \"carol@example.com\", \"password\": \"mypassword\"})\n"
    "        assert res.status_code == 200\n"
    "        assert \"access_token\" in res.json()\n"
    "\n"
    "    def test_login_wrong_password_returns_400(self, client):\n"
    "        client.post(\"/auth/register\", json={\n"
    "            \"email\": \"dave@example.com\", \"password\": \"correct\"})\n"
    "        res = client.post(\"/auth/login\", json={\n"
    "            \"email\": \"dave@example.com\", \"password\": \"wrong\"})\n"
    "        assert res.status_code == 400\n"
    "\n"
    "    def test_login_nonexistent_user_returns_400(self, client):\n"
    "        res = client.post(\"/auth/login\", json={\n"
    "            \"email\": \"ghost@example.com\", \"password\": \"any\"})\n"
    "        assert res.status_code == 400\n"
    "\n"
    "\n"
    "class TestProfile:\n"
    "    def test_get_me_returns_profile(self, client, auth_headers):\n"
    "        res = client.get(\"/auth/me\", headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        data = res.json()\n"
    "        assert data[\"email\"] == \"test@example.com\"\n"
    "        assert data[\"name\"]  == \"Test User\"\n"
    "        assert data[\"city\"]  == \"London\"\n"
    "\n"
    "    def test_update_city_changes_city(self, client, auth_headers):\n"
    "        res = client.put(\"/auth/me\",\n"
    "                         json={\"city\": \"Paris\"},\n"
    "                         headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        assert res.json()[\"city\"] == \"Paris\"\n"
    "\n"
    "    def test_update_to_taken_email_returns_400(self, client, auth_headers):\n"
    "        client.post(\"/auth/register\", json={\n"
    "            \"email\": \"other@example.com\", \"password\": \"pw\"})\n"
    "        res = client.put(\"/auth/me\",\n"
    "                         json={\"email\": \"other@example.com\"},\n"
    "                         headers=auth_headers)\n"
    "        assert res.status_code == 400"
)

heading("G.3  Closet Tests: tests/test_closet.py", 2)
body(
    "The closet test module covers the complete CRUD lifecycle and the multi-parameter "
    "filter behaviour of the list endpoint. The isolation_test verifies that items created "
    "by one user are not visible to another user, which is a critical security property "
    "of the system.",
    indent=0.5
)
cb(
    "# tests/test_closet.py\n"
    "import pytest\n"
    "\n"
    "\n"
    "class TestClosetCRUD:\n"
    "    def test_create_item_returns_201(self, client, auth_headers):\n"
    "        res = client.post(\"/closet\", json={\n"
    "            \"name\": \"Black Jeans\", \"category\": \"Bottom\",\n"
    "            \"color\": \"black\", \"formality\": \"CASUAL\",\n"
    "        }, headers=auth_headers)\n"
    "        assert res.status_code == 201\n"
    "        data = res.json()\n"
    "        assert data[\"name\"]     == \"Black Jeans\"\n"
    "        assert data[\"category\"] == \"Bottom\"\n"
    "        assert data[\"id\"]       > 0\n"
    "\n"
    "    def test_list_items_returns_created_item(self, client, auth_headers, sample_item):\n"
    "        res = client.get(\"/closet\", headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        ids = [i[\"id\"] for i in res.json()]\n"
    "        assert sample_item[\"id\"] in ids\n"
    "\n"
    "    def test_get_item_by_id(self, client, auth_headers, sample_item):\n"
    "        item_id = sample_item[\"id\"]\n"
    "        res = client.get(f\"/closet/{item_id}\", headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        assert res.json()[\"id\"] == item_id\n"
    "\n"
    "    def test_patch_item_updates_color(self, client, auth_headers, sample_item):\n"
    "        item_id = sample_item[\"id\"]\n"
    "        res = client.patch(f\"/closet/{item_id}\",\n"
    "                           json={\"color\": \"blue\"},\n"
    "                           headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        assert res.json()[\"color\"] == \"blue\"\n"
    "        # Name should be unchanged\n"
    "        assert res.json()[\"name\"] == sample_item[\"name\"]\n"
    "\n"
    "    def test_delete_item_returns_204(self, client, auth_headers, sample_item):\n"
    "        item_id = sample_item[\"id\"]\n"
    "        res = client.delete(f\"/closet/{item_id}\", headers=auth_headers)\n"
    "        assert res.status_code == 204\n"
    "        # Verify item is gone\n"
    "        res2 = client.get(\"/closet\", headers=auth_headers)\n"
    "        ids  = [i[\"id\"] for i in res2.json()]\n"
    "        assert item_id not in ids\n"
    "\n"
    "    def test_get_nonexistent_item_returns_404(self, client, auth_headers):\n"
    "        res = client.get(\"/closet/99999\", headers=auth_headers)\n"
    "        assert res.status_code == 404"
)
cb(
    "class TestClosetFilters:\n"
    "    def _seed(self, client, headers):\n"
    "        items = [\n"
    "            {\"name\": \"Navy Blazer\",   \"category\": \"Outerwear\", \"color\": \"navy\",  \"formality\": \"FORMAL\"},\n"
    "            {\"name\": \"White Tee\",     \"category\": \"Top\",       \"color\": \"white\", \"formality\": \"CASUAL\"},\n"
    "            {\"name\": \"Chino Trousers\",\"category\": \"Bottom\",   \"color\": \"beige\", \"formality\": \"MODERATE\"},\n"
    "            {\"name\": \"Chelsea Boots\", \"category\": \"Shoes\",    \"color\": \"black\", \"formality\": \"FORMAL\"},\n"
    "        ]\n"
    "        for item in items:\n"
    "            client.post(\"/closet\", json=item, headers=headers)\n"
    "\n"
    "    def test_category_filter(self, client, auth_headers):\n"
    "        self._seed(client, auth_headers)\n"
    "        res = client.get(\"/closet?category=Top\", headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        items = res.json()\n"
    "        assert len(items) == 1\n"
    "        assert items[0][\"category\"] == \"Top\"\n"
    "\n"
    "    def test_formality_filter(self, client, auth_headers):\n"
    "        self._seed(client, auth_headers)\n"
    "        res = client.get(\"/closet?formality=FORMAL\", headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        for item in res.json():\n"
    "            assert item[\"formality\"] == \"FORMAL\"\n"
    "\n"
    "    def test_search_filter_matches_name(self, client, auth_headers):\n"
    "        self._seed(client, auth_headers)\n"
    "        res = client.get(\"/closet?search=blazer\", headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        assert len(res.json()) == 1\n"
    "        assert \"Blazer\" in res.json()[0][\"name\"]\n"
    "\n"
    "\n"
    "class TestClosetIsolation:\n"
    "    def test_user_cannot_see_other_users_items(self, client, auth_headers):\n"
    "        # User A creates an item\n"
    "        client.post(\"/closet\", json={\n"
    "            \"name\": \"User A Shirt\", \"category\": \"Top\"\n"
    "        }, headers=auth_headers)\n"
    "\n"
    "        # Register User B\n"
    "        client.post(\"/auth/register\", json={\n"
    "            \"email\": \"userb@example.com\", \"password\": \"pw\"})\n"
    "        headers_b = {\"Authorization\": \"Bearer token-userb@example.com\"}\n"
    "\n"
    "        res = client.get(\"/closet\", headers=headers_b)\n"
    "        assert res.status_code == 200\n"
    "        names = [i[\"name\"] for i in res.json()]\n"
    "        assert \"User A Shirt\" not in names"
)

heading("G.4  Outfit Generation Tests: tests/test_outfits.py", 2)
body(
    "The outfit generation tests verify the core recommendation logic: vibe-to-formality "
    "mapping, weather-conditional outerwear inclusion, and the deterministic seeding of "
    "the daily outfit. These tests create a controlled set of closet items and then "
    "assert that the generated outfit contains the expected items based on the request "
    "parameters.",
    indent=0.5
)
cb(
    "# tests/test_outfits.py\n"
    "import pytest\n"
    "\n"
    "\n"
    "@pytest.fixture\n"
    "def full_wardrobe(client, auth_headers):\n"
    "    items = [\n"
    "        {\"name\": \"Formal Shirt\",  \"category\": \"Top\",       \"formality\": \"FORMAL\"},\n"
    "        {\"name\": \"Casual Tee\",    \"category\": \"Top\",       \"formality\": \"CASUAL\"},\n"
    "        {\"name\": \"Dress Trousers\",\"category\": \"Bottom\",   \"formality\": \"FORMAL\"},\n"
    "        {\"name\": \"Jeans\",         \"category\": \"Bottom\",   \"formality\": \"CASUAL\"},\n"
    "        {\"name\": \"Oxford Shoes\",  \"category\": \"Shoes\",    \"formality\": \"FORMAL\"},\n"
    "        {\"name\": \"Trainers\",      \"category\": \"Shoes\",    \"formality\": \"CASUAL\"},\n"
    "        {\"name\": \"Navy Blazer\",   \"category\": \"Outerwear\",\"formality\": \"FORMAL\"},\n"
    "    ]\n"
    "    for item in items:\n"
    "        client.post(\"/closet\", json=item, headers=auth_headers)\n"
    "\n"
    "\n"
    "class TestGenerateOutfit:\n"
    "    def test_generate_returns_outfit_suggestion(self, client, auth_headers, full_wardrobe):\n"
    "        res = client.post(\"/outfits/generate\",\n"
    "                          json={\"context\": \"Office\", \"vibe\": 75},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        data = res.json()\n"
    "        assert \"items\"       in data\n"
    "        assert \"explanation\" in data\n"
    "        assert len(data[\"items\"]) >= 1\n"
    "\n"
    "    def test_high_vibe_selects_formal_items(self, client, auth_headers, full_wardrobe):\n"
    "        res = client.post(\"/outfits/generate\",\n"
    "                          json={\"context\": \"Office\", \"vibe\": 90},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        item_names = [i[\"name\"] for i in res.json()[\"items\"]]\n"
    "        # Formal shirt should be selected over casual tee\n"
    "        assert \"Formal Shirt\" in item_names\n"
    "        assert \"Casual Tee\"   not in item_names\n"
    "\n"
    "    def test_low_vibe_selects_casual_items(self, client, auth_headers, full_wardrobe):\n"
    "        res = client.post(\"/outfits/generate\",\n"
    "                          json={\"context\": \"Casual\", \"vibe\": 10},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        item_names = [i[\"name\"] for i in res.json()[\"items\"]]\n"
    "        assert \"Casual Tee\"   in item_names\n"
    "        assert \"Formal Shirt\" not in item_names\n"
    "\n"
    "    def test_cold_weather_includes_outerwear(self, client, auth_headers, full_wardrobe):\n"
    "        res = client.post(\"/outfits/generate\",\n"
    "                          json={\"context\": \"Travel\",\n"
    "                                \"vibe\": 50,\n"
    "                                \"weather_temp_c\": 5.0},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        categories = [i[\"category\"] for i in res.json()[\"items\"]]\n"
    "        assert \"Outerwear\" in categories\n"
    "\n"
    "    def test_warm_weather_excludes_outerwear(self, client, auth_headers, full_wardrobe):\n"
    "        res = client.post(\"/outfits/generate\",\n"
    "                          json={\"context\": \"Travel\",\n"
    "                                \"vibe\": 50,\n"
    "                                \"weather_temp_c\": 28.0},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        categories = [i[\"category\"] for i in res.json()[\"items\"]]\n"
    "        assert \"Outerwear\" not in categories\n"
    "\n"
    "    def test_empty_closet_returns_empty_items(self, client, auth_headers):\n"
    "        res = client.post(\"/outfits/generate\",\n"
    "                          json={\"context\": \"Office\", \"vibe\": 50},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        assert res.json()[\"items\"] == []"
)
cb(
    "class TestSavedOutfits:\n"
    "    def test_save_and_retrieve_outfit(self, client, auth_headers, full_wardrobe):\n"
    "        gen_res = client.post(\"/outfits/generate\",\n"
    "                              json={\"context\": \"Office\", \"vibe\": 75},\n"
    "                              headers=auth_headers)\n"
    "        generated = gen_res.json()\n"
    "\n"
    "        save_res = client.post(\"/outfits/save\", json={\n"
    "            \"context\":     \"Office\",\n"
    "            \"items\":       generated[\"items\"],\n"
    "            \"explanation\": generated[\"explanation\"],\n"
    "        }, headers=auth_headers)\n"
    "        assert save_res.status_code == 201\n"
    "        saved_id = save_res.json()[\"id\"]\n"
    "\n"
    "        list_res = client.get(\"/outfits/saved\", headers=auth_headers)\n"
    "        assert list_res.status_code == 200\n"
    "        ids = [o[\"id\"] for o in list_res.json()]\n"
    "        assert saved_id in ids\n"
    "\n"
    "    def test_delete_saved_outfit(self, client, auth_headers, full_wardrobe):\n"
    "        gen = client.post(\"/outfits/generate\",\n"
    "                          json={\"context\": \"Office\", \"vibe\": 50},\n"
    "                          headers=auth_headers).json()\n"
    "        saved = client.post(\"/outfits/save\", json={\n"
    "            \"context\": \"Office\", \"items\": gen[\"items\"],\n"
    "            \"explanation\": gen[\"explanation\"]\n"
    "        }, headers=auth_headers).json()\n"
    "\n"
    "        del_res = client.delete(f\"/outfits/saved/{saved['id']}\",\n"
    "                                headers=auth_headers)\n"
    "        assert del_res.status_code == 204\n"
    "\n"
    "        list_res = client.get(\"/outfits/saved\", headers=auth_headers)\n"
    "        ids = [o[\"id\"] for o in list_res.json()]\n"
    "        assert saved[\"id\"] not in ids"
)

heading("G.5  Travel Packing Tests: tests/test_travel.py", 2)
body(
    "The travel tests verify the packing list generation logic. Since the weather fetch "
    "calls the live Open-Meteo API, these tests use pytest's monkeypatch fixture to "
    "substitute a mock _get_weather function that returns controlled temperature values, "
    "ensuring deterministic, network-independent test execution.",
    indent=0.5
)
cb(
    "# tests/test_travel.py\n"
    "import pytest\n"
    "from app.routers import travel as travel_module\n"
    "\n"
    "\n"
    "@pytest.fixture\n"
    "def wardrobe_for_travel(client, auth_headers):\n"
    "    items = [\n"
    "        {\"name\": \"T-Shirt\",       \"category\": \"Top\"},\n"
    "        {\"name\": \"Shorts\",        \"category\": \"Bottom\"},\n"
    "        {\"name\": \"Sandals\",       \"category\": \"Shoes\"},\n"
    "        {\"name\": \"Light Jacket\",  \"category\": \"Outerwear\"},\n"
    "    ]\n"
    "    for item in items:\n"
    "        client.post(\"/closet\", json=item, headers=auth_headers)\n"
    "\n"
    "\n"
    "class TestPackingList:\n"
    "    def test_packing_list_has_correct_day_count(self, client, auth_headers,\n"
    "                                                  wardrobe_for_travel, monkeypatch):\n"
    "        monkeypatch.setattr(travel_module, \"_get_weather\",\n"
    "                            lambda dest: {\"temp_c\": 22.0, \"city\": dest, \"code\": 0})\n"
    "        res = client.post(\"/travel/pack\",\n"
    "                          json={\"destination\": \"Barcelona\", \"days\": 4},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        data = res.json()\n"
    "        assert len(data[\"daily_outfits\"]) == 4\n"
    "\n"
    "    def test_cold_destination_includes_outerwear(self, client, auth_headers,\n"
    "                                                   wardrobe_for_travel, monkeypatch):\n"
    "        monkeypatch.setattr(travel_module, \"_get_weather\",\n"
    "                            lambda dest: {\"temp_c\": 2.0, \"city\": dest, \"code\": 0})\n"
    "        res = client.post(\"/travel/pack\",\n"
    "                          json={\"destination\": \"Reykjavik\", \"days\": 3},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        packing = res.json()[\"packing_list\"]\n"
    "        categories = [i[\"category\"] for i in packing]\n"
    "        assert \"Outerwear\" in categories\n"
    "\n"
    "    def test_warm_destination_excludes_outerwear(self, client, auth_headers,\n"
    "                                                   wardrobe_for_travel, monkeypatch):\n"
    "        monkeypatch.setattr(travel_module, \"_get_weather\",\n"
    "                            lambda dest: {\"temp_c\": 30.0, \"city\": dest, \"code\": 0})\n"
    "        res = client.post(\"/travel/pack\",\n"
    "                          json={\"destination\": \"Dubai\", \"days\": 3},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        packing = res.json()[\"packing_list\"]\n"
    "        categories = [i[\"category\"] for i in packing]\n"
    "        assert \"Outerwear\" not in categories\n"
    "\n"
    "    def test_packing_list_deduplicates_items(self, client, auth_headers,\n"
    "                                              wardrobe_for_travel, monkeypatch):\n"
    "        monkeypatch.setattr(travel_module, \"_get_weather\",\n"
    "                            lambda dest: {\"temp_c\": 22.0, \"city\": dest, \"code\": 0})\n"
    "        res = client.post(\"/travel/pack\",\n"
    "                          json={\"destination\": \"Rome\", \"days\": 5},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        packing = res.json()[\"packing_list\"]\n"
    "        ids = [i[\"id\"] for i in packing]\n"
    "        assert len(ids) == len(set(ids)), \"Duplicate items in packing list\"\n"
    "\n"
    "    def test_empty_closet_returns_400(self, client, auth_headers, monkeypatch):\n"
    "        monkeypatch.setattr(travel_module, \"_get_weather\",\n"
    "                            lambda dest: {\"temp_c\": 20.0, \"city\": dest, \"code\": 0})\n"
    "        res = client.post(\"/travel/pack\",\n"
    "                          json={\"destination\": \"Paris\", \"days\": 3},\n"
    "                          headers=auth_headers)\n"
    "        assert res.status_code == 400"
)

heading("G.6  Wardrobe Gap Analysis Tests: tests/test_suggestions.py", 2)
cb(
    "# tests/test_suggestions.py\n"
    "import pytest\n"
    "\n"
    "\n"
    "class TestGapAnalysis:\n"
    "    def test_empty_closet_returns_build_message(self, client, auth_headers):\n"
    "        res = client.get(\"/suggestions/gaps\", headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        assert \"closet\" in res.json()[\"reason\"].lower()\n"
    "\n"
    "    def test_missing_bottoms_suggests_bottoms(self, client, auth_headers):\n"
    "        for _ in range(5):\n"
    "            client.post(\"/closet\",\n"
    "                        json={\"name\": f\"Top {_}\", \"category\": \"Top\"},\n"
    "                        headers=auth_headers)\n"
    "        res = client.get(\"/suggestions/gaps\", headers=auth_headers)\n"
    "        assert res.status_code == 200\n"
    "        data = res.json()\n"
    "        assert \"Pants\" in data[\"suggestion\"] or \"Jeans\" in data[\"suggestion\"]\n"
    "\n"
    "    def test_missing_shoes_suggests_shoes(self, client, auth_headers):\n"
    "        for cat in [\"Top\", \"Bottom\", \"Outerwear\"]:\n"
    "            for _ in range(2):\n"
    "                client.post(\"/closet\",\n"
    "                            json={\"name\": f\"{cat} {_}\", \"category\": cat},\n"
    "                            headers=auth_headers)\n"
    "        res = client.get(\"/suggestions/gaps\", headers=auth_headers)\n"
    "        data = res.json()\n"
    "        assert \"shoe\" in data[\"suggestion\"].lower() or \\\n"
    "               \"loafer\" in data[\"suggestion\"].lower() or \\\n"
    "               \"sneaker\" in data[\"suggestion\"].lower()"
)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX H — USER INTERFACE WALKTHROUGH
# ══════════════════════════════════════════════════════════════════════════════
pb()
heading("Appendix H: User Interface Walkthrough", 1)
body(
    "This appendix provides a detailed walkthrough of the StylistAI user interface, "
    "describing each screen, its visual design, and the interactions it supports. "
    "The application uses a consistent dark-green colour scheme throughout, with a "
    "deep forest green (#0d1a12) page background, a slightly lighter card background "
    "#121f17, and a vibrant primary green (#22c55e) as the accent colour for buttons, "
    "borders, and interactive elements. This palette was chosen to evoke a premium, "
    "fashion-forward aesthetic while maintaining sufficient contrast for accessibility.",
    indent=0.5
)

heading("H.1  Landing Page", 2)
body(
    "The landing page is the first screen presented to unauthenticated visitors. It "
    "features a full-viewport hero section with an animated gradient headline introducing "
    "the StylistAI concept. Below the hero are three feature highlight cards describing "
    "the Smart Wardrobe, AI Outfit Generator, and Weather Intelligence features, each "
    "with an icon and a two-sentence description. The navigation bar contains a logo on "
    "the left and Sign In / Get Started buttons on the right. The Get Started button "
    "navigates to the registration flow; Sign In navigates to the login form. The page "
    "is fully responsive: on mobile, the hero text scales down and the feature cards "
    "stack vertically.",
    indent=0.5
)

heading("H.2  Authentication Flow", 2)
body(
    "The authentication flow uses a single AuthFlow component that renders either a "
    "Login form or a Sign Up form depending on the initialView prop. Both forms share "
    "the same visual design: a centred card with the StylistAI logo, a title, labelled "
    "input fields with real-time validation feedback, and a submit button. The Sign Up "
    "form includes Name and City fields in addition to Email and Password. Validation "
    "errors from the backend (duplicate email, wrong password) are displayed as red "
    "inline error messages beneath the relevant field. The forms switch between login and "
    "signup modes via a text link at the bottom of the card ('Already have an account? "
    "Sign In'), preserving any entered email address to avoid re-typing.",
    indent=0.5
)

heading("H.3  Dashboard View", 2)
body(
    "After login, the user is directed to the Wardrobe view (first visit) or, on "
    "subsequent visits, to the Dashboard. The Dashboard occupies the main content area "
    "to the right of the fixed sidebar. The top section displays a personalised greeting "
    "('Good morning, Samagan.') with the current date, followed by a row of context "
    "chips showing the current weather temperature, the user's city, the next calendar "
    "event title, and the total number of closet items. Below the context chips is the "
    "main outfit card, a two-panel layout: the left panel shows a 2x2 grid of garment "
    "thumbnail images, and the right panel shows the outfit name (item names joined by "
    "dots), the AI-generated explanation in italic, and three action buttons: "
    "'Wear Today' (logs the outfit), 'Refresh' (fetches a different combination), and "
    "'Customise' (navigates to the Generator). At the bottom of the dashboard, a "
    "Wardrobe Analysis card displays the current gap analysis recommendation.",
    indent=0.5
)

heading("H.4  Wardrobe View", 2)
body(
    "The Wardrobe View is the most feature-rich screen in the application. The top "
    "section contains a filter bar with four controls: a category dropdown (All, Top, "
    "Bottom, Outerwear, Shoes, Accessory), a colour selector (rendered as a row of "
    "circular colour swatches), a formality filter (CASUAL, MODERATE, FORMAL, "
    "UNIVERSAL), and a free-text search box. Below the filter bar is the upload panel, "
    "a dashed-border dropzone that accepts both click-to-browse and drag-and-drop file "
    "uploads alongside metadata fields for Name, Category, Colour, and Formality. "
    "The main area renders the filtered closet items as a responsive grid of cards. "
    "Each card is approximately 200 x 240 pixels and shows the garment image with a "
    "transparent background on a dark surface, the item name, a category badge, and a "
    "colour swatch. Clicking a card opens an edit modal with pre-populated fields and "
    "a Delete button. On mobile, the grid reduces to two columns and the filter bar "
    "collapses to a single 'Filters' button that opens a bottom drawer.",
    indent=0.5
)

heading("H.5  Style Generator View", 2)
body(
    "The Generator View guides the user through a two-step outfit creation flow. In Step "
    "1 (Config), the user selects an occasion from a 2x3 grid of occasion tiles, each "
    "containing an icon, label, and one-line description. The selected tile is "
    "highlighted with a primary-green border. Below the occasion grid, a vibe slider "
    "labelled 'Comfort to Style' allows the user to adjust the formality of the "
    "generated outfit on a 0-100 scale. The slider label updates dynamically as the "
    "user drags it, showing descriptors like 'Pure Comfort', 'Balanced', or "
    "'Style Focus'. A weather chip shows the current temperature if the user's city is "
    "configured, providing context for the temperature-aware recommendation. The "
    "Generate button at the bottom of the config panel triggers the API call and "
    "transitions to a loading state. Step 2 (Result) shows the generated outfit as a "
    "list of item cards with category labels and formality badges, the explanation text, "
    "and three action buttons: 'Save', 'Virtual Try-On', and 'Generate Again'. The "
    "Virtual Try-On button sends the outfit to POST /outfits/try-on and displays the "
    "1024x1024 DALL-E 3 image in a full-screen modal overlay.",
    indent=0.5
)

heading("H.6  Travel Packing View", 2)
body(
    "The Travel Packing View presents a simple form with two inputs: a Destination text "
    "field (city or country name) and a Days counter (min 1, max 30). After submitting "
    "the form, the view transitions to a results layout showing a weather banner at the "
    "top (e.g. 'Expected weather in Paris: 14.2°C'), followed by the master packing "
    "list as a checklist and a day-by-day accordion. Each day section shows a row of "
    "outfit item thumbnails with their names and categories. Users can tap individual "
    "packing list items to check them off as they pack, with a visual strikethrough and "
    "reduced opacity applied to checked items. The check state is managed in local "
    "component state and is not persisted to the server.",
    indent=0.5
)

heading("H.7  Sidebar Navigation", 2)
body(
    "The sidebar is a fixed left panel on desktop screens (width >= 768 px). It contains "
    "the StylistAI logo at the top, five navigation items (Dashboard, Wardrobe, "
    "Generator, Travel, Profile), a dark mode toggle switch at the bottom, and a logout "
    "button. Each navigation item consists of an icon, a label, and an active indicator "
    "(a left-border highlight in the primary green colour). On mobile screens, the "
    "sidebar is replaced by a bottom navigation bar showing only the icons, with labels "
    "appearing as tooltips on long press. The sidebar state (open/collapsed) is "
    "controlled by the activeTab prop passed from App.jsx and automatically syncs with "
    "the browser URL via the useEffect hook in App.jsx.",
    indent=0.5
)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX I — TECHNOLOGY STACK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
pb()
heading("Appendix I: Technology Stack Reference", 1)
body(
    "This appendix provides a comprehensive reference of all technologies, libraries, "
    "and external services used in the StylistAI project, including version numbers as "
    "used in the final submission and the rationale for selecting each dependency.",
    indent=0.5
)

heading("I.1  Backend Dependencies", 2)
para()
t_back = doc.add_table(rows=1, cols=4)
t_back.style = "Table Grid"
for cell, txt in zip(t_back.rows[0].cells, ["Package", "Version", "Purpose", "Rationale"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)
for row in [
    ("fastapi",        "0.115.x", "ASGI web framework",                "Automatic OpenAPI docs, Pydantic integration, async-first design"),
    ("uvicorn",        "0.32.x",  "ASGI server",                       "Recommended by FastAPI; supports hot-reload in development"),
    ("sqlalchemy",     "2.0.x",   "ORM and query layer",               "Database-agnostic; clean Declarative Base API; strong type annotations"),
    ("pydantic",       "v2",      "Data validation and serialisation",  "Native FastAPI integration; fast Rust-based validation core"),
    ("bcrypt",         "4.x",     "Password hashing",                  "Industry-standard adaptive cost function; resists brute-force attacks"),
    ("python-dotenv",  "1.x",     "Environment variable loading",      "Standard .env file support; keeps secrets out of source code"),
    ("openai",         "1.x",     "DALL-E 3 virtual try-on",           "Official SDK for OpenAI API; handles auth and rate-limit retries"),
    ("rembg",          "2.x",     "Background removal (U2-Net)",       "Convenient wrapper around the U2-Net ONNX model for background segmentation"),
    ("pillow",         "10.x",    "Image format support for rembg",    "Required by rembg for PNG output with alpha channel"),
    ("httpx",          "0.27.x",  "HTTP client for weather API",       "Modern async-capable alternative to requests; used in weather router"),
    ("pytest",         "8.x",     "Test runner",                       "Industry-standard Python testing framework; extensive plugin ecosystem"),
]:
    r = t_back.add_row()
    for i, val in enumerate(row):
        r.cells[i].text = val
        for p in r.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(9)
para()

heading("I.2  Frontend Dependencies", 2)
para()
t_front = doc.add_table(rows=1, cols=4)
t_front.style = "Table Grid"
for cell, txt in zip(t_front.rows[0].cells, ["Package", "Version", "Purpose", "Rationale"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)
for row in [
    ("react",             "19.x",   "UI component framework",            "Concurrent mode features; stable hooks API; largest ecosystem"),
    ("react-dom",         "19.x",   "DOM rendering layer",               "Paired with React for browser environments"),
    ("react-router-dom",  "7.x",    "Client-side routing",               "Nested route support; loader/action patterns; v7 stable API"),
    ("tailwindcss",       "3.x",    "Utility-first CSS framework",        "No runtime overhead; PurgeCSS integration removes unused styles"),
    ("lucide-react",      "0.x",    "SVG icon library",                  "Tree-shakeable; consistent stroke-based design; 1400+ icons"),
    ("@google/material-design-icons", "—", "Material Symbols icon font", "Provides additional icons not in Lucide (checkroom, auto_awesome)"),
    ("vite",              "5.x",    "Build tool and dev server",         "Sub-second HMR; ES module-based; far faster than CRA webpack"),
    ("framer-motion",     "11.x",   "Animation library",                 "Declarative spring animations for card transitions and page fades"),
]:
    r = t_front.add_row()
    for i, val in enumerate(row):
        r.cells[i].text = val
        for p in r.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(9)
para()

heading("I.3  External APIs and Services", 2)
para()
t_api = doc.add_table(rows=1, cols=4)
t_api.style = "Table Grid"
for cell, txt in zip(t_api.rows[0].cells, ["Service", "Endpoint Used", "Purpose", "Cost / Auth"]):
    cell.text = txt
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)
for row in [
    ("Open-Meteo",     "api.open-meteo.com/v1/forecast",               "Real-time weather temperature and weather code",          "Free; no API key required"),
    ("Open-Meteo Geo", "geocoding-api.open-meteo.com/v1/search",       "City name to latitude/longitude geocoding",               "Free; no API key required"),
    ("OpenAI DALL-E 3","api.openai.com/v1/images/generations",          "1024x1024 outfit portrait generation",                   "Paid; OPENAI_API_KEY env var"),
]:
    r = t_api.add_row()
    for i, val in enumerate(row):
        r.cells[i].text = val
        for p in r.cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(9)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(DOCX)
print(f"Saved: {DOCX}")

from docx import Document as D2
d = D2(DOCX)
words = sum(len(p.text.split()) for p in d.paragraphs if p.text.strip())
tables_words = sum(
    len(cell.text.split())
    for table in d.tables
    for row in table.rows
    for cell in row.cells
    if cell.text.strip()
)
total = words + tables_words
print(f"Paragraph words:  {words}")
print(f"Table words:      {tables_words}")
print(f"Total words:      {total}")
print(f"Approx pages:     ~{total // 350}")
